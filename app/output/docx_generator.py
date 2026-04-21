"""
DOCX Marketing Plan Generator — Adaptive version.
Walks agent output JSON trees and renders them regardless of key naming conventions.
"""
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _make_field_run(p_element, field_code: str):
    """Insert a Word field using fldChar BEGIN/SEPARATE/END pattern.

    This format is compatible with Microsoft Word, Apple Pages, Google Docs,
    and LibreOffice — unlike the simpler fldSimple approach.
    """
    def _rpr():
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8pt
        rpr.append(sz)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "999999")
        rpr.append(color)
        return rpr

    # BEGIN
    r_begin = OxmlElement("w:r")
    r_begin.append(_rpr())
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fld_begin)
    p_element.append(r_begin)

    # INSTRUCTION
    r_instr = OxmlElement("w:r")
    r_instr.append(_rpr())
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {field_code} "
    r_instr.append(instr_text)
    p_element.append(r_instr)

    # SEPARATE
    r_sep = OxmlElement("w:r")
    r_sep.append(_rpr())
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fld_sep)
    p_element.append(r_sep)

    # DEFAULT TEXT (shown before fields are updated)
    r_default = OxmlElement("w:r")
    r_default.append(_rpr())
    t = OxmlElement("w:t")
    t.text = "1"
    r_default.append(t)
    p_element.append(r_default)

    # END
    r_end = OxmlElement("w:r")
    r_end.append(_rpr())
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end.append(fld_end)
    p_element.append(r_end)


def _add_page_numbers(doc: Document):
    """Add centered 'Page X of Y' to the document footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = p.add_run("Page ")
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        _make_field_run(p._element, "PAGE")

        run2 = p.add_run(" of ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        _make_field_run(p._element, "NUMPAGES")


# ─── Agent display order and section titles ────────────────
AGENT_SECTIONS = [
    ("chief_strategist", "Executive Summary & Grand Strategy"),
    ("intake_analyst", "Product Assessment"),
    ("behavioral_scientist", "Behavioral Framework"),
    ("psychometrics_expert", "Audience Segmentation & Psychometrics"),
    ("competitive_intelligence", "Competitive Landscape"),
    ("social_strategist", "Social Media Strategy"),
    ("creative_director", "Creative Brief & Campaign Deliverables"),
    ("stakeholder_agent", "Stakeholder Interview Framework"),
]

# Keys to skip (metadata, not content)
SKIP_KEYS = {
    "productName", "product_name", "analysisVersion", "analysis_version",
    "product_bible_entries", "psychometricAnalysisVersion", "productAnalyzed",
    "expertName", "expert_name", "analysisDate", "analysis_date",
    "preparedBy", "prepared_by", "version", "grandStrategyTitle",
    "behavioralScientist", "missing_information_alert",
    "sources_cited", "sources",  # rendered in bibliography section
    "research_questions",        # rendered separately via _render_research_questions
    "phase_2_disclaimer",        # rendered separately as an advisory callout
}


def _safe_str(text) -> str:
    """Strip NULL bytes and XML-illegal control characters from a string."""
    import re
    s = str(text)
    # Remove NULL bytes and control chars (except tab, newline, carriage return)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)


def _humanize_key(key: str) -> str:
    """Convert camelCase or snake_case key to readable heading."""
    import re
    # Handle camelCase
    s = re.sub(r'([a-z])([A-Z])', r'\1 \2', key)
    # Handle snake_case
    s = s.replace('_', ' ')
    # Remove common prefixes like part1_, part2_
    s = re.sub(r'^part\s*\d+\s*', '', s, flags=re.IGNORECASE)
    # Title case
    s = s.strip().title()
    return s


def _render_value(doc, value, depth=2):
    """Recursively render a JSON value into the document."""
    if value is None:
        return

    if isinstance(value, str):
        if len(value) > 0:
            doc.add_paragraph(_safe_str(value))

    elif isinstance(value, bool):
        doc.add_paragraph(_safe_str(value))

    elif isinstance(value, (int, float)):
        doc.add_paragraph(_safe_str(value))

    elif isinstance(value, list):
        for item in value:
            if isinstance(item, str):
                doc.add_paragraph(_safe_str(item), style="List Bullet")
            elif isinstance(item, dict):
                _render_dict(doc, item, depth=depth)
            elif isinstance(item, list):
                for sub in item:
                    if isinstance(sub, str):
                        doc.add_paragraph(_safe_str(sub), style="List Bullet")
                    else:
                        _render_value(doc, sub, depth=depth)

    elif isinstance(value, dict):
        _render_dict(doc, value, depth=depth)


def _render_dict(doc, data: dict, depth=2):
    """Render a dictionary. Tries to be smart about structure."""
    if not data:
        return

    # Check if this looks like a simple key-value record (all string/number values)
    simple_values = {}
    complex_values = {}
    for k, v in data.items():
        if k in SKIP_KEYS:
            continue
        if isinstance(v, (str, int, float, bool)) and v != "":
            simple_values[k] = v
        elif isinstance(v, (list, dict)) and v:
            complex_values[k] = v

    # Render simple values as bold-label paragraphs
    for k, v in simple_values.items():
        label = _humanize_key(k)
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(_safe_str(v))

    # Render complex values with headings
    heading_level = min(depth, 4)  # Cap at Heading 4
    for k, v in complex_values.items():
        label = _humanize_key(k)
        if heading_level <= 3:
            doc.add_heading(label, level=heading_level)
        else:
            p = doc.add_paragraph()
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(11)
        _render_value(doc, v, depth=heading_level + 1)


def _collect_all_citations(agent_outputs: dict) -> list[dict]:
    """Gather and deduplicate citations from all agent outputs."""
    seen_urls = set()
    citations = []
    for agent_data in agent_outputs.values():
        if not isinstance(agent_data, dict):
            continue
        # Structured format (new)
        for src in agent_data.get("sources_cited", []):
            if isinstance(src, dict):
                url = src.get("url", "")
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    citations.append(src)
        # Legacy format: plain URL strings
        for url in agent_data.get("sources", []):
            if isinstance(url, str) and url not in seen_urls:
                seen_urls.add(url)
                citations.append({
                    "url": url,
                    "title": url.split("//")[-1].split("/")[0],
                    "description": "",
                    "finding": "",
                })
    return citations


def _render_bibliography(doc, citations: list[dict]):
    """Render annotated bibliography entries into the DOCX.

    Updated 2026-04-15 to surface the fields the human reviewer needs to
    verify a citation: authors, article title, publication, year, URL,
    MLA line, and a confidence signal flagging fabrication risk.
    """
    # Confidence colors — the user should see at a glance which citations to trust.
    CONF_COLOR = {
        "verified":          RGBColor(0x27, 0xAE, 0x60),  # green
        "likely":            RGBColor(0xF3, 0x9C, 0x12),  # amber — verify before quoting
        "general-consensus": RGBColor(0x88, 0x88, 0x88),  # gray — not a specific paper
    }

    for i, cite in enumerate(citations, 1):
        # Prefer new fields; fall back to legacy (`title`) for older runs.
        article_title = cite.get("article_title") or cite.get("title") or "Untitled Source"
        authors       = cite.get("authors", "")
        publication   = cite.get("publication") or cite.get("journal", "")
        year          = cite.get("year", "")
        url           = cite.get("url", "")
        doi           = cite.get("doi", "")
        citation_mla  = cite.get("citation_mla", "")
        description   = cite.get("description", "")
        finding       = cite.get("finding", "")
        relevance     = cite.get("relevance", "")
        source_type   = cite.get("source_type", "")
        confidence    = cite.get("confidence", "")

        # Line 1: bold title (article title is the key verifiable field)
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {article_title}")
        run.bold = True
        run.font.size = Pt(10)

        # Confidence tag inline with the title — the user's first honesty signal.
        if confidence:
            tag = p.add_run(f"   [{confidence.upper()}]")
            tag.font.size = Pt(8)
            tag.bold = True
            tag.font.color.rgb = CONF_COLOR.get(confidence, RGBColor(0x88, 0x88, 0x88))

        # Line 2: authors + publication + year (the verifiable trio)
        byline_parts = []
        if authors:
            byline_parts.append(authors)
        if publication:
            byline_parts.append(publication)
        if year:
            byline_parts.append(str(year))
        if source_type:
            byline_parts.append(source_type)
        if byline_parts:
            p_by = doc.add_paragraph()
            run_by = p_by.add_run(" · ".join(byline_parts))
            run_by.font.size = Pt(9)
            run_by.italic = True
            run_by.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Line 3: MLA line the user can copy-paste straight into a bibliography.
        if citation_mla:
            p_mla = doc.add_paragraph()
            lbl = p_mla.add_run("MLA: ")
            lbl.bold = True
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            body = p_mla.add_run(citation_mla)
            body.font.size = Pt(9)

        # URL — shown as plain text (python-docx does not natively render
        # clickable hyperlinks without more plumbing, but the URL is still
        # copy-pasteable and visually distinct).
        if url and url != "N/A":
            p_url = doc.add_paragraph()
            run_url = p_url.add_run(url)
            run_url.font.size = Pt(8)
            run_url.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
            run_url.underline = True

        if doi:
            p_doi = doc.add_paragraph()
            run_doi = p_doi.add_run(f"DOI: {doi}")
            run_doi.font.size = Pt(8)
            run_doi.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

        # Description / finding / relevance — the contextual block.
        for label, body_text in [("Contains", description), ("Key finding", finding), ("Relevance", relevance)]:
            if body_text:
                p_x = doc.add_paragraph()
                lbl = p_x.add_run(f"{label}: ")
                lbl.italic = True
                lbl.font.size = Pt(9)
                lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                body = p_x.add_run(_safe_str(body_text))
                body.font.size = Pt(9)

        # Spacing between entries
        if i < len(citations):
            doc.add_paragraph("")


def _render_phase_2_disclaimer(doc, disclaimer_text: str):
    """Visually-distinct callout separating Phase 1 research from Phase 2 synthesis."""
    if not disclaimer_text:
        disclaimer_text = (
            "The synthesis that follows (framework, tactics, segments, recommendations) is only "
            "as sound as the Phase 1 research and citations above. Verify those first — especially "
            "any citation marked [LIKELY] or [GENERAL-CONSENSUS] — before relying on the strategic "
            "conclusions drawn from them."
        )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    lbl = p.add_run("⚠ Phase 1 → Phase 2 handoff. ")
    lbl.bold = True
    lbl.font.size = Pt(10)
    lbl.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    body = p.add_run(_safe_str(disclaimer_text))
    body.font.size = Pt(10)
    body.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    body.italic = True
    doc.add_paragraph("")


def _render_review_advisory(doc):
    """Top-of-document advisory explaining the two-phase verification workflow.

    Added 2026-04-15 in response to human-in-loop feedback. Readers should know
    up front that citations must be verified before the synthesis layer is trusted.
    """
    doc.add_heading("How to read this report", level=2)
    advisory = (
        "This report is produced in two phases. Phase 1 is research: sources, findings, and "
        "direct answers to specific research questions. Phase 2 is synthesis: the strategic "
        "framework, tactics, and recommendations that follow from that research.\n\n"
        "Before trusting Phase 2, spot-check Phase 1. Every citation carries a confidence tag:\n"
        "  • VERIFIED — the paper exists exactly as cited; safe to rely on.\n"
        "  • LIKELY — the authors and topic are real; the title may be paraphrased. Search authors + topic to find the real paper.\n"
        "  • GENERAL-CONSENSUS — no specific paper; field-level finding. Treat as opinion informed by the field.\n\n"
        "If any citation cannot be found via its MLA line or URL, flag the whole section — the "
        "synthesis built on it is suspect. This is not a weakness of this tool; it is the "
        "two-phase workflow working as intended."
    )
    for paragraph in advisory.split("\n\n"):
        p = doc.add_paragraph()
        r = p.add_run(paragraph)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph("")


def _render_research_questions(doc, questions: list[dict]):
    """Render the Wayfinders-style per-question research block.

    Each question gets: question → data points → summary → application →
    caveats → marketing implication → per-question source list.
    The human reviewer asked for this format explicitly — they want to see
    each question answered with its sources attached, not a single giant
    bibliography at the end.
    """
    if not questions:
        return

    doc.add_heading("Research Questions & Findings", level=2)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Each research question below is answered with multiple data points, a summary, "
        "application to this product, caveats, a marketing implication, and its own source list. "
        "Verify the sources at the end of each question before trusting the synthesis that follows."
    )
    intro_run.italic = True
    intro_run.font.size = Pt(10)
    intro_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for i, q in enumerate(questions, 1):
        # Question heading
        question = q.get("question", "Untitled question")
        doc.add_heading(f"Q{i}. {question}", level=3)

        # Data points
        data_points = q.get("data_points") or []
        if data_points:
            dp_lbl = doc.add_paragraph()
            r = dp_lbl.add_run("Data points")
            r.bold = True
            r.font.size = Pt(10)

            for dp in data_points:
                finding = dp.get("finding", "")
                mechanism = dp.get("mechanism") or dp.get("framework_basis", "")
                strength = dp.get("evidence_strength", "")
                source_ref = dp.get("source_ref", "")

                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(_safe_str(finding))
                run.font.size = Pt(10)

                meta_bits = []
                if mechanism:
                    meta_bits.append(f"Mechanism: {mechanism}")
                if strength:
                    meta_bits.append(f"Evidence: {strength}")
                if source_ref:
                    meta_bits.append(f"Source: {source_ref}")
                if meta_bits:
                    meta_run = p.add_run(f"  ({' · '.join(meta_bits)})")
                    meta_run.font.size = Pt(8)
                    meta_run.italic = True
                    meta_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        # Summary / application / caveats / marketing implication
        for label, key in [
            ("Summary", "summary"),
            ("Application to this product", "application"),
            ("Caveats & limitations", "caveats_and_limitations"),
            ("Marketing implication", "marketing_implication"),
        ]:
            val = q.get(key)
            if val:
                p = doc.add_paragraph()
                lbl = p.add_run(f"{label}: ")
                lbl.bold = True
                lbl.font.size = Pt(10)
                body = p.add_run(_safe_str(val))
                body.font.size = Pt(10)

        # Per-question source list
        qs = q.get("question_sources") or []
        if qs:
            p = doc.add_paragraph()
            lbl = p.add_run("Sources for this question: ")
            lbl.bold = True
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
            body = p.add_run("; ".join(_safe_str(s) for s in qs))
            body.font.size = Pt(9)
            body.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

        doc.add_paragraph("")


def generate_marketing_plan(
    project_name: str,
    project_type: str,
    agent_outputs: dict[str, dict],
    output_path: str,
    citations: list[dict] = None,
    stakeholder_questions: list[dict] = None,
) -> str:
    """Generate the comprehensive marketing plan DOCX from all agent outputs.

    Stakeholder section positioning is dynamic:
    - No answers yet → section appears FIRST (action item for reader)
    - Any answers exist → section appears LAST with full Q&A audit trail
    - No questions at all → section omitted
    """
    doc = Document()
    _setup_styles(doc)

    sq = stakeholder_questions or []
    has_any_questions = len(sq) > 0
    has_any_answers = any(q.get("answer") for q in sq)

    # Sections excluding stakeholder (we'll insert it manually)
    main_sections = [(k, t) for k, t in AGENT_SECTIONS if k != "stakeholder_agent"]

    # ─── Cover Page ─────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project_name.upper())
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Comprehensive Marketing Strategy")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x6C, 0x6C, 0x80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run = meta.add_run(f"\nProduct Type: {project_type.replace('_', ' ').title()}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run = meta.add_run("\nPowered by Idideration")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    cost_line = meta.add_run("\n\nBehavioral Science \u2022 Neuroscience \u2022 Psychometrics \u2022 Game Theory")
    cost_line.font.size = Pt(9)
    cost_line.font.color.rgb = RGBColor(0x99, 0x99, 0xBB)

    doc.add_page_break()

    # ─── Table of Contents ──────────────────────────────────
    doc.add_heading("Table of Contents", level=1)

    toc_sections = []
    if has_any_questions and not has_any_answers:
        # Stakeholder questions first
        toc_sections.append("Stakeholder Interview Questions")
        toc_sections.extend(t for _, t in main_sections)
    else:
        toc_sections.extend(t for _, t in main_sections)
        if has_any_questions:
            toc_sections.append("Stakeholder Interviews — Q&A Record")
    toc_sections.append("Annotated Bibliography")

    for i, title_text in enumerate(toc_sections, 1):
        doc.add_paragraph(f"{i}. {title_text}")
    doc.add_page_break()

    # ─── Review advisory (how to read this report) ──────────
    # Added 2026-04-15: tells the reader about the two-phase verify-then-trust flow
    # and the citation confidence tags. User feedback: they need to know up front
    # that they should verify sources before relying on synthesis.
    _render_review_advisory(doc)
    doc.add_page_break()

    # Build numbered map so rendered sections match TOC numbering
    _section_numbers = {t: i for i, t in enumerate(toc_sections, 1)}

    # ─── Helper: render stakeholder section ─────────────────
    def _render_stakeholder_section(heading: str):
        doc.add_heading(heading, level=1)
        if not sq:
            doc.add_paragraph("(No stakeholder questions generated.)")
            doc.add_page_break()
            return

        answered = [q for q in sq if q.get("answer")]
        outstanding = [q for q in sq if not q.get("answer")]

        current_role = None

        if not has_any_answers:
            # Pure question list — grouped by role, with answer lines
            for q in sq:
                role = (q.get("target_role") or "General").replace("_", " ").title()
                if role != current_role:
                    current_role = role
                    doc.add_heading(role, level=2)
                doc.add_paragraph(q.get("question", ""), style="List Number")
                if q.get("purpose"):
                    p = doc.add_paragraph(f"  Purpose: {q['purpose']}")
                    p.runs[0].italic = True
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph("  Answer: _______________________________________________")
                doc.add_paragraph("")
        else:
            # Q&A audit trail
            if answered:
                doc.add_heading("Answered", level=2)
                current_role = None
                for q in answered:
                    role = (q.get("target_role") or "General").replace("_", " ").title()
                    if role != current_role:
                        current_role = role
                        p = doc.add_paragraph(role)
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(10)
                        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x66)

                    p_q = doc.add_paragraph()
                    p_q.add_run("Q: ").bold = True
                    p_q.add_run(q.get("question", ""))

                    p_a = doc.add_paragraph()
                    run_a = p_a.add_run("A: ")
                    run_a.bold = True
                    run_a.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
                    answered_by = q.get("answered_by", "")
                    prefix = f"({answered_by}) " if answered_by else ""
                    p_a.add_run(f"{prefix}{q.get('answer', '')}")
                    doc.add_paragraph("")

            if outstanding:
                doc.add_heading("Outstanding", level=2)
                current_role = None
                for q in outstanding:
                    role = (q.get("target_role") or "General").replace("_", " ").title()
                    if role != current_role:
                        current_role = role
                        p = doc.add_paragraph(role)
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(10)
                        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x66)

                    p_q = doc.add_paragraph()
                    run_label = p_q.add_run("Q: ")
                    run_label.bold = True
                    p_q.add_run(q.get("question", ""))

                    p_out = doc.add_paragraph()
                    run_out = p_out.add_run("  [ OUTSTANDING ]")
                    run_out.italic = True
                    run_out.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
                    run_out.font.size = Pt(9)
                    doc.add_paragraph("")

        doc.add_page_break()

    # ─── Render sections ────────────────────────────────────
    def _numbered(title: str) -> str:
        n = _section_numbers.get(title)
        return f"{n}. {title}" if n else title

    if has_any_questions and not has_any_answers:
        _render_stakeholder_section(_numbered("Stakeholder Interview Questions"))

    for agent_key, section_title in main_sections:
        agent_data = agent_outputs.get(agent_key, {})
        doc.add_heading(_numbered(section_title), level=1)
        if not agent_data:
            doc.add_paragraph("(No data from this agent.)")
        else:
            # Phase 1: per-question research block (new, preferred reading order)
            research_questions = agent_data.get("research_questions") or []
            if research_questions:
                _render_research_questions(doc, research_questions)

            # Phase 1 → Phase 2 handoff callout
            if research_questions or agent_data.get("phase_2_disclaimer"):
                _render_phase_2_disclaimer(doc, agent_data.get("phase_2_disclaimer", ""))

            # Phase 2: everything else (framework, tactics, segments, etc.)
            filtered = {k: v for k, v in agent_data.items() if k not in SKIP_KEYS and v}
            _render_dict(doc, filtered, depth=2)
        doc.add_page_break()

    if has_any_questions and has_any_answers:
        _render_stakeholder_section(_numbered("Stakeholder Interviews — Q&A Record"))

    # ─── Annotated Bibliography ─────────────────────────────
    all_citations = _collect_all_citations(agent_outputs)
    if all_citations:
        doc.add_heading(_numbered("Annotated Bibliography"), level=1)
        _render_bibliography(doc, all_citations)
        doc.add_page_break()

    # ─── Footer note ────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by Idideration \u2022 idideration.com")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    _add_page_numbers(doc)
    doc.save(output_path)
    return output_path


def generate_single_agent_docx(
    project_name: str,
    agent_key: str,
    agent_output: dict,
    output_path: str,
) -> str:
    """Generate a DOCX for a single agent's output."""
    doc = Document()
    _setup_styles(doc)

    # Look up section title from AGENT_SECTIONS
    section_title = None
    for key, title in AGENT_SECTIONS:
        if key == agent_key:
            section_title = title
            break
    if not section_title:
        section_title = _humanize_key(agent_key)

    doc.add_heading(f"{project_name} — {section_title}", level=1)

    meta = doc.add_paragraph()
    run = meta.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.add_paragraph("")

    # Phase 1: per-question research block
    research_questions = agent_output.get("research_questions") or []
    if research_questions:
        _render_research_questions(doc, research_questions)
    if research_questions or agent_output.get("phase_2_disclaimer"):
        _render_phase_2_disclaimer(doc, agent_output.get("phase_2_disclaimer", ""))

    filtered = {k: v for k, v in agent_output.items() if k not in SKIP_KEYS and v}
    if filtered:
        _render_dict(doc, filtered, depth=2)
    elif not research_questions:
        doc.add_paragraph("(No data from this agent.)")

    # ─── Bibliography for this agent ────────────────────────
    sources = agent_output.get("sources_cited", [])
    legacy = agent_output.get("sources", [])
    combined = list(sources)
    for url in legacy:
        if isinstance(url, str):
            combined.append({
                "url": url,
                "title": url.split("//")[-1].split("/")[0],
                "description": "",
                "finding": "",
            })
    if combined:
        doc.add_page_break()
        doc.add_heading("Annotated Bibliography", level=1)
        _render_bibliography(doc, combined)

    _add_page_numbers(doc)
    doc.save(output_path)
    return output_path


def generate_research_briefs_docx(
    project_name: str,
    briefs: list[dict],
    output_path: str,
) -> str:
    """Generate a DOCX containing one or more research briefs.

    Each brief is a dict with keys: question, output_json, cost_usd, created_at.
    """
    doc = Document()
    _setup_styles(doc)

    # ─── Cover Page ─────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project_name.upper())
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Research Briefs")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x6C, 0x6C, 0x80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    meta.add_run(f"\n{len(briefs)} brief{'s' if len(briefs) != 1 else ''}")

    doc.add_page_break()

    # ─── Table of Contents ──────────────────────────────────
    if len(briefs) > 1:
        doc.add_heading("Contents", level=1)
        for i, b in enumerate(briefs, 1):
            q = b.get("question", "Untitled")
            p = doc.add_paragraph(f"{i}. {q[:120]}{'...' if len(q) > 120 else ''}")
            p.style = doc.styles["List Number"]
        doc.add_page_break()

    # ─── Render each brief ──────────────────────────────────
    for i, b in enumerate(briefs, 1):
        question = b.get("question", "Untitled")
        out = b.get("output_json") or {}
        created = b.get("created_at", "")
        cost = b.get("cost_usd", 0)

        # Section header
        doc.add_heading(f"{'Brief ' + str(i) + ': ' if len(briefs) > 1 else ''}{question}", level=1)

        # Meta line
        meta_parts = []
        if created:
            meta_parts.append(str(created)[:16])
        if cost:
            meta_parts.append(f"${cost:.4f}")
        if meta_parts:
            p = doc.add_paragraph(" · ".join(meta_parts))
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            p.runs[0].italic = True

        if not out:
            doc.add_paragraph("(No output available.)")
            doc.add_page_break()
            continue

        doc.add_paragraph("")

        # Summary
        if out.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(out["summary"])

        # Key Findings
        if out.get("key_findings"):
            doc.add_heading("Key Findings", level=2)
            for f in out["key_findings"]:
                if not isinstance(f, dict):
                    continue
                conf = f.get("confidence", "").upper()
                finding = f.get("finding", "")
                explanation = f.get("explanation", "")
                source_hint = f.get("source_hint", "")
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"[{conf}] " if conf else "")
                run.bold = True
                run.font.color.rgb = (
                    RGBColor(0x16, 0xa3, 0x4a) if conf == "HIGH" else
                    RGBColor(0xd9, 0x77, 0x06) if conf == "MEDIUM" else
                    RGBColor(0x99, 0x99, 0x99)
                )
                p.add_run(finding)
                if explanation:
                    p2 = doc.add_paragraph(f"    {explanation}")
                    p2.runs[0].font.size = Pt(9)
                    p2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                if source_hint:
                    p3 = doc.add_paragraph(f"    Source: {source_hint}")
                    p3.runs[0].font.size = Pt(9)
                    p3.runs[0].italic = True
                    p3.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Academic & Industry perspectives
        if out.get("academic_perspective"):
            doc.add_heading("Academic Perspective", level=2)
            doc.add_paragraph(out["academic_perspective"])

        if out.get("industry_perspective"):
            doc.add_heading("Industry Perspective", level=2)
            doc.add_paragraph(out["industry_perspective"])

        # Data Points
        if out.get("data_points"):
            doc.add_heading("Data Points", level=2)
            for dp in out["data_points"]:
                if not isinstance(dp, dict):
                    continue
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(dp.get("stat", ""))
                run.bold = True
                if dp.get("source"):
                    p.add_run(f" — {dp['source']}")

        # Marketing Implications
        if out.get("implications_for_marketing"):
            doc.add_heading("Marketing Implications", level=2)
            doc.add_paragraph(out["implications_for_marketing"])

        # Caveats
        if out.get("caveats_and_limitations"):
            doc.add_heading("Caveats & Limitations", level=2)
            doc.add_paragraph(out["caveats_and_limitations"])

        # Related Topics
        if out.get("related_topics"):
            doc.add_heading("Related Topics", level=2)
            for t in out["related_topics"]:
                doc.add_paragraph(str(t), style="List Bullet")

        # Sources
        sources = out.get("sources_cited", [])
        if sources:
            doc.add_heading("Sources", level=2)
            _render_bibliography(doc, [s for s in sources if isinstance(s, dict)])

        if i < len(briefs):
            doc.add_page_break()

    _add_page_numbers(doc)
    doc.save(output_path)
    return output_path


def _setup_styles(doc: Document):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Ensure List Bullet 2 exists
    if "List Bullet 2" not in doc.styles:
        doc.styles.add_style("List Bullet 2", WD_STYLE_TYPE.PARAGRAPH)


def generate_stakeholder_questions_docx(
    project_name: str,
    questions: list[dict],
    output_path: str,
) -> str:
    """Generate a stakeholder interview questions document."""
    doc = Document()
    _setup_styles(doc)

    doc.add_heading(f"{project_name} \u2014 Stakeholder Interview Questions", level=1)
    doc.add_paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    current_role = None
    for q in questions:
        role = q.get("target_role", "general")
        if role != current_role:
            current_role = role
            doc.add_heading(role.replace("_", " ").title(), level=2)

        doc.add_paragraph(q.get("question", ""), style="List Number")
        if q.get("purpose"):
            p = doc.add_paragraph(f"  Purpose: {q['purpose']}")
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph("  Answer: _______________________________________________")
        doc.add_paragraph("")

    _add_page_numbers(doc)
    doc.save(output_path)
    return output_path
