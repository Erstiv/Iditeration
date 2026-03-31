"""Project management routes."""
import json
import logging
from fastapi import APIRouter, Depends, Request, Form, BackgroundTasks, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from pathlib import Path
from datetime import datetime, timezone

from app.database import get_db
from app.models import (
    Project, ProjectType, CrewRun, AgentRun, AgentName,
    BibleEntry, BibleScope, BibleCategory, StakeholderQuestion,
    OutputDocument, RunStatus, AgentNote, ResearchBrief,
)
from app.crews.marketing_bible import MarketingBibleTool
from app.output.docx_generator import generate_marketing_plan, generate_stakeholder_questions_docx, generate_single_agent_docx
from app.output.html_renderer import render_agent_output_html, render_agent_output_editable_html
from app.crews.crew_runner import create_crew_run, execute_crew_run, get_crew_run_status, DEFAULT_AGENT_ORDER, rerun_single_agent, continue_crew_run, AGENT_CLASS_MAP, process_stakeholder_answers
from app.crews.agents.research_brief_agent import run_research_brief
from app.config import PROJECTS_DIR, GEMINI_API_KEY

logger = logging.getLogger(__name__)

router = APIRouter()

# ─── Agent interaction helpers ──────────────────────────────

AGENT_DEFAULT_BIBLE_CATEGORY = {
    "research_agent":           "research_findings",
    "intake_analyst":           "product_overview",
    "behavioral_scientist":     "audience_segments",
    "psychometrics_expert":     "audience_segments",
    "competitive_intelligence": "competitive_data",
    "social_strategist":        "social_data",
    "chief_strategist":         "strategy",
    "creative_director":        "creative",
    "stakeholder_agent":        "stakeholder_input",
}

PUSH_BIBLE_CATEGORIES = [
    ("research_findings",  "Research Findings"),
    ("product_overview",   "Product Overview"),
    ("audience_segments",  "Audience Segments"),
    ("competitive_data",   "Competitive Data"),
    ("social_data",        "Social Data"),
    ("strategy",           "Strategy"),
    ("creative",           "Creative"),
    ("stakeholder_input",  "Stakeholder Input"),
]


def _estimate_note_cost(model: str, input_tokens: int, output_tokens: int) -> float:
    if "pro" in model:
        return round((input_tokens / 1_000_000) * 1.25 + (output_tokens / 1_000_000) * 10.00, 4)
    return round((input_tokens / 1_000_000) * 0.15 + (output_tokens / 1_000_000) * 0.60, 4)
templates = Jinja2Templates(directory=Path(__file__).parent.parent / "templates")


# ─── Project CRUD ───────────────────────────────────────────

@router.get("/", response_class=HTMLResponse)
async def dashboard(request: Request, db: Session = Depends(get_db)):
    projects = db.query(Project).order_by(Project.updated_at.desc()).all()
    return templates.TemplateResponse(request, "dashboard.html", {
        "projects": projects,
        "project_types": [t.value for t in ProjectType],
    })


@router.get("/projects/new", response_class=HTMLResponse)
async def new_project_form(request: Request):
    return templates.TemplateResponse(request, "new_project.html", {
        "project_types": [t.value for t in ProjectType],
    })


@router.post("/projects/create")
async def create_project(
    background_tasks: BackgroundTasks,
    name: str = Form(...),
    project_type: str = Form(...),
    description: str = Form(""),
    raw_data: str = Form(""),
    narralytica_show_id: int = Form(None),
    # New intake form fields
    url: str = Form(""),
    elevator_pitch: str = Form(""),
    key_people: str = Form(""),
    target_audience: str = Form(""),
    known_competitors: str = Form(""),
    social_accounts: str = Form(""),
    distribution: str = Form(""),
    # Research options
    auto_research: str = Form(None),
    generate_questions: str = Form(None),  # Legacy, kept for backward compat
    full_pipeline: str = Form(None),
    db: Session = Depends(get_db),
):
    # Compile all intake data into raw_data if the new form fields are provided
    extra_sections = []
    if url:
        extra_sections.append(f"URL: {url}")
    if elevator_pitch:
        extra_sections.append(f"Elevator Pitch: {elevator_pitch}")
    if key_people:
        extra_sections.append(f"Key People: {key_people}")
    if target_audience:
        extra_sections.append(f"Target Audience: {target_audience}")
    if known_competitors:
        extra_sections.append(f"Known Competitors: {known_competitors}")
    if social_accounts:
        extra_sections.append(f"Social Accounts: {social_accounts}")
    if distribution:
        extra_sections.append(f"Distribution: {distribution}")

    # Merge extra fields into raw_data
    combined_raw = raw_data or ""
    if extra_sections:
        intake_block = "\n--- Intake Form Data ---\n" + "\n".join(extra_sections)
        combined_raw = (combined_raw + "\n\n" + intake_block).strip() if combined_raw else intake_block

    project = Project(
        user_id=1,  # Single-user for now
        name=name,
        project_type=ProjectType(project_type),
        description=description,
        raw_data=combined_raw,
        narralytica_show_id=narralytica_show_id,
    )
    db.add(project)
    db.commit()

    # Create project output directory
    project_dir = PROJECTS_DIR / str(project.id)
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "output").mkdir(exist_ok=True)

    # Handle auto-research and pipeline options
    if full_pipeline:
        # Run research agent + full pipeline (all agents in order)
        agents = DEFAULT_AGENT_ORDER  # Already includes RESEARCH_AGENT first
        crew_run = create_crew_run(db, project.id, agents)
        background_tasks.add_task(execute_crew_run, db, crew_run.id)
        return RedirectResponse(url=f"/projects/{project.id}/runs/{crew_run.id}", status_code=303)
    elif auto_research:
        # Run just the research agent (and optionally stakeholder)
        agents = [AgentName.RESEARCH_AGENT]
        if generate_questions:
            agents.append(AgentName.STAKEHOLDER_AGENT)
        crew_run = create_crew_run(db, project.id, agents)
        background_tasks.add_task(execute_crew_run, db, crew_run.id)
        return RedirectResponse(url=f"/projects/{project.id}/runs/{crew_run.id}", status_code=303)
    elif generate_questions:
        # Run just the stakeholder agent
        agents = [AgentName.STAKEHOLDER_AGENT]
        crew_run = create_crew_run(db, project.id, agents)
        background_tasks.add_task(execute_crew_run, db, crew_run.id)
        return RedirectResponse(url=f"/projects/{project.id}/runs/{crew_run.id}", status_code=303)

    return RedirectResponse(url=f"/projects/{project.id}", status_code=303)


@router.get("/projects/{project_id}", response_class=HTMLResponse)
async def project_detail(request: Request, project_id: int, db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        return RedirectResponse(url="/", status_code=303)

    crew_runs = db.query(CrewRun).filter(CrewRun.project_id == project_id).order_by(CrewRun.created_at.desc()).all()
    bible = MarketingBibleTool(db, project_id)
    product_entries = bible.get_all_project_entries()
    stakeholder_qs = db.query(StakeholderQuestion).filter(
        StakeholderQuestion.project_id == project_id
    ).all()
    output_docs = db.query(OutputDocument).filter(
        OutputDocument.project_id == project_id
    ).order_by(OutputDocument.generated_at.desc()).all()
    research_briefs = db.query(ResearchBrief).filter(
        ResearchBrief.project_id == project_id
    ).order_by(ResearchBrief.created_at.desc()).all()

    return templates.TemplateResponse(request, "project.html", {
        "project": project,
        "crew_runs": crew_runs,
        "product_entries": product_entries,
        "stakeholder_questions": stakeholder_qs,
        "output_docs": output_docs,
        "research_briefs": research_briefs,
        "agent_names": [a.value for a in AgentName],
    })


@router.post("/projects/{project_id}/delete")
async def delete_project(project_id: int, db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        return RedirectResponse(url="/", status_code=303)
    db.delete(project)
    db.commit()
    return RedirectResponse(url="/", status_code=303)


@router.post("/projects/{project_id}/update")
async def update_project(
    project_id: int,
    name: str = Form(None),
    description: str = Form(None),
    raw_data: str = Form(None),
    db: Session = Depends(get_db),
):
    project = db.query(Project).filter(Project.id == project_id).first()
    if name:
        project.name = name
    if description is not None:
        project.description = description
    if raw_data is not None:
        project.raw_data = raw_data
    db.commit()
    return RedirectResponse(url=f"/projects/{project_id}", status_code=303)


# ─── Crew Runs ──────────────────────────────────────────────

@router.post("/projects/{project_id}/run")
async def start_crew_run(
    project_id: int,
    background_tasks: BackgroundTasks,
    agents: list[str] = Form(None),
    db: Session = Depends(get_db),
):
    if agents:
        agent_list = [AgentName(a) for a in agents]
    else:
        agent_list = DEFAULT_AGENT_ORDER

    crew_run = create_crew_run(db, project_id, agent_list)
    background_tasks.add_task(execute_crew_run, db, crew_run.id)
    return RedirectResponse(url=f"/projects/{project_id}/runs/{crew_run.id}", status_code=303)


@router.get("/projects/{project_id}/runs/{run_id}", response_class=HTMLResponse)
async def run_status_page(request: Request, project_id: int, run_id: int, db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == project_id).first()
    status = get_crew_run_status(db, run_id)
    bible = MarketingBibleTool(db, project_id)
    active_directives = bible.get_project_entries(category=BibleCategory.DIRECTIVES)
    return templates.TemplateResponse(request, "run_status.html", {
        "project": project,
        "run_status": status,
        "active_directives": active_directives,
    })


@router.get("/api/projects/{project_id}/runs/{run_id}/status")
async def run_status_api(project_id: int, run_id: int, db: Session = Depends(get_db)):
    """API endpoint for polling run status (used by HTMX/JS)."""
    return get_crew_run_status(db, run_id)


# ─── Output Generation ─────────────────────────────────────

@router.post("/projects/{project_id}/generate-docx")
async def generate_docx(project_id: int, run_id: int = Form(...), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == project_id).first()
    crew_run = db.query(CrewRun).filter(CrewRun.id == run_id).first()

    if not crew_run or crew_run.status != RunStatus.COMPLETED:
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)

    # Gather all agent outputs
    agent_outputs = {}
    for ar in crew_run.agent_runs:
        if ar.output_json:
            agent_outputs[ar.agent_name.value] = ar.output_json

    # Generate DOCX
    output_dir = PROJECTS_DIR / str(project_id) / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{project.name.replace(' ', '_')}_Marketing_Plan_{timestamp}.docx"
    output_path = str(output_dir / filename)

    generate_marketing_plan(
        project_name=project.name,
        project_type=project.project_type.value,
        agent_outputs=agent_outputs,
        output_path=output_path,
    )

    # Record in database
    doc_record = OutputDocument(
        project_id=project_id,
        crew_run_id=run_id,
        doc_type="marketing_plan",
        file_path=output_path,
        file_name=filename,
    )
    db.add(doc_record)
    db.commit()

    return RedirectResponse(url=f"/projects/{project_id}", status_code=303)


@router.get("/projects/{project_id}/download/{doc_id}")
async def download_document(project_id: int, doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(OutputDocument).filter(OutputDocument.id == doc_id).first()
    if not doc or not Path(doc.file_path).exists():
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)
    return FileResponse(doc.file_path, filename=doc.file_name)


# ─── Single-Agent Operations ──────────────────────────────

@router.get("/api/projects/{project_id}/runs/{run_id}/agents/{agent_name}/output", response_class=HTMLResponse)
async def agent_output_fragment(
    request: Request,
    project_id: int,
    run_id: int,
    agent_name: str,
    db: Session = Depends(get_db),
):
    """Return rendered HTML fragment of a single agent's output."""
    agent_run = db.query(AgentRun).filter(
        AgentRun.crew_run_id == run_id,
        AgentRun.agent_name == AgentName(agent_name),
    ).first()

    if not agent_run or not agent_run.output_json:
        rendered_html = '<p style="color:var(--text-dim);">No output available.</p>'
    else:
        rendered_html = render_agent_output_html(agent_run.output_json)

    return templates.TemplateResponse(request, "partials/_agent_output.html", {
        "rendered_html": rendered_html,
    })


@router.get("/api/projects/{project_id}/runs/{run_id}/agents/{agent_name}/edit", response_class=HTMLResponse)
async def agent_edit_fragment(
    request: Request,
    project_id: int,
    run_id: int,
    agent_name: str,
    db: Session = Depends(get_db),
):
    """Return editable HTML form of a single agent's output."""
    agent_run = db.query(AgentRun).filter(
        AgentRun.crew_run_id == run_id,
        AgentRun.agent_name == AgentName(agent_name),
    ).first()

    if not agent_run or not agent_run.output_json:
        edit_html = '<p style="color:var(--text-dim);">No output available to edit.</p>'
    else:
        edit_html = render_agent_output_editable_html(agent_run.output_json)

    form_html = (
        f'<form hx-post="/api/projects/{project_id}/runs/{run_id}/agents/{agent_name}/save-edit" '
        f'hx-target="#output-{agent_name}" hx-swap="innerHTML">'
        f'{edit_html}'
        f'<div style="margin-top:1rem;display:flex;gap:0.5rem;">'
        f'<button type="submit" class="btn btn-primary btn-sm">Save Edits</button>'
        f'<button type="button" class="btn btn-outline btn-sm" '
        f'onclick="htmx.ajax(\'GET\',\'/api/projects/{project_id}/runs/{run_id}/agents/{agent_name}/output\','
        f'{{target:\'#output-{agent_name}\',swap:\'innerHTML\'}})">Cancel</button>'
        f'</div></form>'
    )
    return HTMLResponse(
        f'<div style="padding:1rem;background:var(--bg);border:1px solid var(--accent);'
        f'border-radius:6px;margin-top:0.5rem;max-height:600px;overflow-y:auto;font-size:0.85rem;">'
        f'{form_html}</div>'
    )


@router.post("/api/projects/{project_id}/runs/{run_id}/agents/{agent_name}/save-edit", response_class=HTMLResponse)
async def agent_save_edit(
    request: Request,
    project_id: int,
    run_id: int,
    agent_name: str,
    db: Session = Depends(get_db),
):
    """Save edited agent output fields back to the database."""
    agent_run = db.query(AgentRun).filter(
        AgentRun.crew_run_id == run_id,
        AgentRun.agent_name == AgentName(agent_name),
    ).first()

    if not agent_run or not agent_run.output_json:
        return HTMLResponse('<p style="color:var(--red);">Agent run not found.</p>')

    form_data = await request.form()
    output = agent_run.output_json.copy() if isinstance(agent_run.output_json, dict) else {}

    # Reconstruct JSON from form field paths (e.g., "platform_audit.0.platform")
    for field_name, value in form_data.items():
        _set_nested(output, field_name.split("."), str(value))

    agent_run.output_json = output
    db.commit()

    rendered_html = render_agent_output_html(output)
    return templates.TemplateResponse(request, "partials/_agent_output.html", {
        "rendered_html": rendered_html,
    })


def _set_nested(obj, keys: list[str], value):
    """Set a value in a nested dict/list structure using a list of path keys."""
    for i, key in enumerate(keys[:-1]):
        next_key = keys[i + 1]
        if key.isdigit():
            key = int(key)
            while len(obj) <= key:
                obj.append({} if not next_key.isdigit() else [])
            obj = obj[key]
        else:
            if key not in obj:
                obj[key] = [] if next_key.isdigit() else {}
            obj = obj[key]

    final_key = keys[-1]
    if final_key.isdigit():
        idx = int(final_key)
        while len(obj) <= idx:
            obj.append("")
        obj[idx] = value
    else:
        obj[final_key] = value


@router.post("/projects/{project_id}/runs/{run_id}/agents/{agent_name}/export-docx")
async def export_agent_docx(
    project_id: int,
    run_id: int,
    agent_name: str,
    db: Session = Depends(get_db),
):
    """Generate and download a DOCX for a single agent's output."""
    project = db.query(Project).filter(Project.id == project_id).first()
    agent_run = db.query(AgentRun).filter(
        AgentRun.crew_run_id == run_id,
        AgentRun.agent_name == AgentName(agent_name),
    ).first()

    if not agent_run or not agent_run.output_json:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    output_dir = PROJECTS_DIR / str(project_id) / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    agent_label = agent_name.replace("_", " ").title().replace(" ", "_")
    filename = f"{project.name.replace(' ', '_')}_{agent_label}_{timestamp}.docx"
    output_path = str(output_dir / filename)

    generate_single_agent_docx(
        project_name=project.name,
        agent_key=agent_name,
        agent_output=agent_run.output_json,
        output_path=output_path,
    )

    # Record in database
    doc_record = OutputDocument(
        project_id=project_id,
        crew_run_id=run_id,
        doc_type="agent_export",
        file_path=output_path,
        file_name=filename,
    )
    db.add(doc_record)
    db.commit()

    return FileResponse(output_path, filename=filename)


@router.post("/projects/{project_id}/runs/{run_id}/continue")
async def continue_run(
    project_id: int,
    run_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    """Continue a crew run from the first pending agent."""
    crew_run = db.query(CrewRun).filter(CrewRun.id == run_id).first()
    if not crew_run:
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)
    running = any(ar.status == RunStatus.RUNNING for ar in crew_run.agent_runs)
    if running:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)
    background_tasks.add_task(continue_crew_run, db, run_id)
    return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)


@router.post("/projects/{project_id}/runs/{run_id}/agents/{agent_name}/rerun")
async def rerun_agent(
    project_id: int,
    run_id: int,
    agent_name: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    """Rerun a single agent within an existing crew run."""
    # Check no agent is currently running in this crew run
    crew_run = db.query(CrewRun).filter(CrewRun.id == run_id).first()
    if not crew_run:
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)

    running = any(ar.status == RunStatus.RUNNING for ar in crew_run.agent_runs)
    if running:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    try:
        agent_enum = AgentName(agent_name)
    except ValueError:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    background_tasks.add_task(rerun_single_agent, db, run_id, agent_enum)
    return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)


# ─── Agent Interaction (Ask / Refocus) ───────────────────────

@router.get("/api/projects/{project_id}/runs/{run_id}/agents/{agent_name}/chat", response_class=HTMLResponse)
async def agent_chat_fragment(
    request: Request,
    project_id: int,
    run_id: int,
    agent_name: str,
    db: Session = Depends(get_db),
):
    """Load the Ask/Refocus panel for a single agent."""
    agent_run = db.query(AgentRun).filter(
        AgentRun.crew_run_id == run_id,
        AgentRun.agent_name == AgentName(agent_name),
    ).first()
    existing_notes = agent_run.notes if agent_run else []
    return templates.TemplateResponse(request, "partials/_agent_chat.html", {
        "existing_notes": existing_notes,
        "project_id": project_id,
        "run_id": run_id,
        "agent_name": agent_name,
        "bible_categories": PUSH_BIBLE_CATEGORIES,
        "default_category": AGENT_DEFAULT_BIBLE_CATEGORY.get(agent_name, "research_findings"),
    })


@router.post("/projects/{project_id}/runs/{run_id}/agents/{agent_name}/ask", response_class=HTMLResponse)
async def agent_ask(
    request: Request,
    project_id: int,
    run_id: int,
    agent_name: str,
    question: str = Form(...),
    db: Session = Depends(get_db),
):
    """Ask a follow-up question to an agent about its output."""
    from google import genai as _genai
    from google.genai import types as _genai_types

    agent_run = db.query(AgentRun).filter(
        AgentRun.crew_run_id == run_id,
        AgentRun.agent_name == AgentName(agent_name),
    ).first()

    if not agent_run or not agent_run.output_json:
        return HTMLResponse('<p style="color:var(--red);">Agent has no output to ask about.</p>')

    model = agent_run.model_used or "gemini-2.5-flash"
    agent_class = AGENT_CLASS_MAP.get(AgentName(agent_name))
    system_prompt = agent_class.system_prompt if agent_class else ""

    ask_prompt = (
        f"=== YOUR PREVIOUS OUTPUT (CONTEXT) ===\n"
        f"{json.dumps(agent_run.output_json, indent=2, default=str)[:20000]}\n\n"
        f"=== USER QUESTION ===\n{question}\n\n"
        f"Answer the question directly based on your output above. "
        f"Respond in plain text, not JSON. Be specific and concise."
    )

    try:
        client = _genai.Client(api_key=GEMINI_API_KEY)
        response = client.models.generate_content(
            model=model,
            contents=ask_prompt,
            config=_genai_types.GenerateContentConfig(
                system_instruction=system_prompt,
                temperature=0.3,
                max_output_tokens=4096,
            ),
        )
        answer = response.text or "(No response)"
        in_tok = out_tok = 0
        if hasattr(response, "usage_metadata") and response.usage_metadata:
            in_tok = getattr(response.usage_metadata, "prompt_token_count", 0) or 0
            out_tok = getattr(response.usage_metadata, "candidates_token_count", 0) or 0
        cost = _estimate_note_cost(model, in_tok, out_tok)
    except Exception as e:
        answer = f"(Error getting response: {e})"
        in_tok = out_tok = cost = 0

    note = AgentNote(
        crew_run_id=run_id,
        agent_run_id=agent_run.id,
        project_id=project_id,
        question=question,
        answer=answer,
        model_used=model,
        input_tokens=in_tok,
        output_tokens=out_tok,
        cost_usd=cost,
    )
    db.add(note)
    db.commit()
    db.refresh(note)

    return templates.TemplateResponse(request, "partials/_note_item.html", {
        "note": note,
        "project_id": project_id,
        "bible_categories": PUSH_BIBLE_CATEGORIES,
        "default_category": AGENT_DEFAULT_BIBLE_CATEGORY.get(agent_name, "research_findings"),
    })


@router.post("/projects/{project_id}/notes/{note_id}/push-to-bible", response_class=HTMLResponse)
async def push_note_to_bible(
    request: Request,
    project_id: int,
    note_id: int,
    category: str = Form(...),
    title: str = Form(...),
    db: Session = Depends(get_db),
):
    """Push an AgentNote's answer into the Product Bible."""
    note = db.query(AgentNote).filter(
        AgentNote.id == note_id,
        AgentNote.project_id == project_id,
    ).first()

    if not note or note.bible_entry_id:
        return HTMLResponse('<span class="badge badge-completed">In Bible</span>')

    bible = MarketingBibleTool(db, project_id)
    entry = bible.add_entry(
        category=BibleCategory(category),
        title=title,
        content=note.answer,
        source=f"ask:{note.agent_run.agent_name.value}",
    )
    note.bible_entry_id = entry.id
    db.commit()

    return HTMLResponse('<span class="badge badge-completed">In Bible</span>')


@router.post("/projects/{project_id}/runs/{run_id}/agents/{agent_name}/refocus")
async def agent_refocus(
    project_id: int,
    run_id: int,
    agent_name: str,
    background_tasks: BackgroundTasks,
    directive: str = Form(...),
    db: Session = Depends(get_db),
):
    """Add a directive to the Product Bible and rerun the agent."""
    crew_run = db.query(CrewRun).filter(CrewRun.id == run_id).first()
    if not crew_run:
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)

    running = any(ar.status == RunStatus.RUNNING for ar in crew_run.agent_runs)
    if running:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    try:
        agent_enum = AgentName(agent_name)
    except ValueError:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    bible = MarketingBibleTool(db, project_id)
    bible.add_entry(
        category=BibleCategory.DIRECTIVES,
        title=f"[{agent_name.replace('_', ' ').title()}] {directive[:80]}",
        content=directive,
        source=f"refocus:{agent_name}",
    )

    background_tasks.add_task(rerun_single_agent, db, run_id, agent_enum)
    return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)


@router.post("/projects/{project_id}/runs/{run_id}/agents/{agent_name}/rerun-guided")
async def agent_rerun_guided(
    project_id: int,
    run_id: int,
    agent_name: str,
    background_tasks: BackgroundTasks,
    guidance: str = Form(...),
    db: Session = Depends(get_db),
):
    """Rerun an agent with one-time guidance (not persisted to Bible)."""
    crew_run = db.query(CrewRun).filter(CrewRun.id == run_id).first()
    if not crew_run:
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)

    running = any(ar.status == RunStatus.RUNNING for ar in crew_run.agent_runs)
    if running:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    try:
        agent_enum = AgentName(agent_name)
    except ValueError:
        return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)

    background_tasks.add_task(rerun_single_agent, db, run_id, agent_enum, guidance=guidance)
    return RedirectResponse(url=f"/projects/{project_id}/runs/{run_id}", status_code=303)


# ─── Knowledge Base ─────────────────────────────────────────

@router.post("/projects/{project_id}/bible/add")
async def add_bible_entry(
    project_id: int,
    category: str = Form(...),
    title: str = Form(...),
    content: str = Form(...),
    db: Session = Depends(get_db),
):
    bible = MarketingBibleTool(db, project_id)
    bible.add_entry(
        category=BibleCategory(category),
        title=title,
        content=content,
        source="manual",
    )
    return RedirectResponse(url=f"/projects/{project_id}", status_code=303)


# ─── File Upload & AI Parsing ──────────────────────────────

@router.post("/api/parse-files")
async def parse_files(files: list[UploadFile] = File(...)):
    """Accept uploaded PDF/DOCX/TXT files, extract text, parse with Gemini."""
    all_text_parts = []

    for f in files:
        filename = (f.filename or "").lower()
        raw_bytes = await f.read()

        try:
            if filename.endswith(".pdf"):
                import PyPDF2
                import io
                reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
                pages_text = [page.extract_text() or "" for page in reader.pages]
                all_text_parts.append(f"--- {f.filename} ---\n" + "\n".join(pages_text))

            elif filename.endswith((".docx", ".doc")):
                import docx
                import io
                doc = docx.Document(io.BytesIO(raw_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                all_text_parts.append(f"--- {f.filename} ---\n" + "\n".join(paragraphs))

            elif filename.endswith(".txt"):
                text = raw_bytes.decode("utf-8", errors="replace")
                all_text_parts.append(f"--- {f.filename} ---\n" + text)

            else:
                all_text_parts.append(f"--- {f.filename} (unsupported format, skipped) ---")
        except Exception as e:
            logger.error(f"Error extracting text from {f.filename}: {e}")
            all_text_parts.append(f"--- {f.filename} (extraction error: {e}) ---")

    combined_text = "\n\n".join(all_text_parts)

    if not combined_text.strip():
        return JSONResponse({"error": "No text could be extracted from the uploaded files.", "raw_text": ""})

    # Try Gemini parsing
    parsed = {
        "elevator_pitch": "",
        "key_people": "",
        "target_audience": "",
        "known_competitors": "",
        "social_accounts": "",
        "distribution": "",
        "raw_notes": "",
    }

    prompt = f"""You are parsing unstructured marketing/product notes into structured fields.
Given the following text extracted from uploaded documents, extract and organize
the information into these categories. Return ONLY valid JSON, no markdown fences.

{{
  "elevator_pitch": "A concise pitch for the product if found",
  "key_people": "Key people, creators, cast, team members mentioned",
  "target_audience": "Any audience or demographic information",
  "known_competitors": "Any competitors or comparable products mentioned",
  "social_accounts": "Any social media accounts, handles, or URLs",
  "distribution": "Distribution channels, platforms, or release info",
  "raw_notes": "Any remaining information that doesn't fit the above categories"
}}

Leave fields as empty strings if no relevant information was found.

=== EXTRACTED TEXT ===
{combined_text}"""

    try:
        from google import genai
        from google.genai import types as genai_types
        client = genai.Client(api_key=GEMINI_API_KEY)
        logger.info(f"Parsing {len(files)} files, {len(combined_text)} chars of text...")
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                temperature=0.2,
                max_output_tokens=8192,
                response_mime_type="application/json",
            ),
        )
        response_text = (response.text or "").strip()
        logger.info(f"Gemini response length: {len(response_text)}")
        if response_text:
            # Strip markdown code fences if present
            import re
            response_text = re.sub(r"^```(?:json)?\s*\n?", "", response_text)
            response_text = re.sub(r"\n?```\s*$", "", response_text).strip()
            parsed = json.loads(response_text)
        else:
            logger.warning("Gemini returned empty response for file parsing")
            parsed["raw_notes"] = combined_text
    except Exception as e:
        logger.error(f"Gemini parsing failed: {type(e).__name__}: {e}")
        # Fallback: return raw text so nothing is lost
        parsed["raw_notes"] = combined_text

    # Always include the raw extracted text
    parsed["raw_text"] = combined_text
    return JSONResponse(parsed)


# ─── Stakeholder Questions ──────────────────────────────────

@router.post("/projects/{project_id}/stakeholder/answer")
async def answer_stakeholder_question(
    project_id: int,
    question_id: int = Form(...),
    answer: str = Form(...),
    answered_by: str = Form(""),
    db: Session = Depends(get_db),
):
    sq = db.query(StakeholderQuestion).filter(StakeholderQuestion.id == question_id).first()
    if sq:
        sq.answer = answer
        sq.answered_by = answered_by
        sq.answered_at = datetime.now(timezone.utc)
        db.commit()
    return RedirectResponse(url=f"/projects/{project_id}", status_code=303)


@router.post("/projects/{project_id}/stakeholder/process")
async def process_stakeholder(
    project_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    """Process answered stakeholder questions — writes findings to Product Bible."""
    answered = db.query(StakeholderQuestion).filter(
        StakeholderQuestion.project_id == project_id,
        StakeholderQuestion.answer.isnot(None),
    ).count()
    if answered == 0:
        return RedirectResponse(url=f"/projects/{project_id}", status_code=303)

    background_tasks.add_task(process_stakeholder_answers, db, project_id)
    return RedirectResponse(url=f"/projects/{project_id}", status_code=303)


# ─── Research Briefs ─────────────────────────────────────────

@router.post("/projects/{project_id}/research-brief")
async def create_research_brief(
    project_id: int,
    background_tasks: BackgroundTasks,
    question: str = Form(...),
    db: Session = Depends(get_db),
):
    """Start a new research brief query."""
    brief = ResearchBrief(project_id=project_id, question=question.strip())
    db.add(brief)
    db.commit()
    db.refresh(brief)
    background_tasks.add_task(run_research_brief, db, brief.id)
    return RedirectResponse(url=f"/projects/{project_id}", status_code=303)


@router.get("/api/projects/{project_id}/research-briefs/{brief_id}", response_class=HTMLResponse)
async def get_research_brief(
    request: Request,
    project_id: int,
    brief_id: int,
    db: Session = Depends(get_db),
):
    """Return HTML fragment for a research brief result."""
    brief = db.query(ResearchBrief).filter(
        ResearchBrief.id == brief_id,
        ResearchBrief.project_id == project_id,
    ).first()
    if not brief:
        return HTMLResponse('<p class="text-dim text-sm">Brief not found.</p>')
    return templates.TemplateResponse(request, "partials/_research_brief_result.html", {
        "brief": brief,
        "project_id": project_id,
    })


@router.post("/api/projects/{project_id}/research-briefs/{brief_id}/save-to-bible",
             response_class=HTMLResponse)
async def save_brief_to_bible(
    request: Request,
    project_id: int,
    brief_id: int,
    category: str = Form("research_findings"),
    db: Session = Depends(get_db),
):
    """Save a research brief's key findings to the Product Bible."""
    brief = db.query(ResearchBrief).filter(
        ResearchBrief.id == brief_id,
        ResearchBrief.project_id == project_id,
    ).first()
    if not brief or not brief.output_json:
        return HTMLResponse('<span class="badge badge-failed">Not ready</span>')

    output = brief.output_json
    question = output.get("question", brief.question)
    summary = output.get("summary", "")
    implications = output.get("implications_for_marketing", "")

    content = summary
    if implications:
        content += f"\n\n### Marketing Implications\n{implications}"

    # Include key findings if present
    findings = output.get("key_findings", [])
    if findings:
        content += "\n\n### Key Findings\n"
        for f in findings[:5]:
            if isinstance(f, dict):
                conf = f.get("confidence", "")
                content += f"- [{conf.upper()}] {f.get('finding', '')}\n"

    bible = MarketingBibleTool(db, project_id)
    try:
        cat = BibleCategory(category)
    except ValueError:
        cat = BibleCategory.RESEARCH_FINDINGS

    bible.add_entry(
        category=cat,
        title=f"Research Brief: {question[:80]}",
        content=content.strip(),
        source="research_brief",
    )
    return HTMLResponse('<span class="badge badge-completed">Saved to Bible ✓</span>')
