#!/usr/bin/env python3
"""Generate the Idideration Executive Summary DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = "/Users/JERS/Idideration/projects/Idideration_Executive_Summary.docx"

doc = Document()

# -- Global defaults --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = Pt(12)

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

DARK_BLUE = RGBColor(0x1A, 0x1A, 0x2E)
GRAY_666 = RGBColor(0x66, 0x66, 0x66)
GRAY_888 = RGBColor(0x88, 0x88, 0x88)
GRAY_999 = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_BG = "F2F2F2"


def add_centered(text, size, bold=False, italic=False, color=DARK_BLUE, space_after=Pt(1)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(size + 2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)


def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(15)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    return p


def add_body(text, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_body_with_bold_label(bold_text, normal_text, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(12)
    r1 = p.add_run(bold_text)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    r2 = p.add_run(normal_text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(
        '<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex)
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, size=Pt(9), bold=False, color=RGBColor(0x33, 0x33, 0x33)):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color


# ============================================================
# PAGE 1
# ============================================================

# Title block
add_centered("IDIDERATION", 28, bold=True, color=DARK_BLUE, space_after=Pt(2))
add_centered("AI-Powered Marketing Strategy Platform", 16, color=DARK_BLUE, space_after=Pt(2))
add_centered("Replacing Marketing Departments with Intelligence", 12, italic=True, color=GRAY_666, space_after=Pt(1))
add_centered("March 2026", 10, color=GRAY_999, space_after=Pt(2))

add_hr()

# The Problem
add_section_header("The Problem")
add_body(
    "Traditional marketing agencies charge $15-50K per month, take weeks to deliver strategy, "
    "and rarely incorporate behavioral science. Most marketing plans are built on gut instinct, "
    "not neuroscience. Small studios, indie creators, and startups simply cannot access CMO-level "
    "strategic thinking. The gap between what academic research tells us about human behavior and "
    "how marketing is actually practiced is enormous."
)

# The Solution
add_section_header("The Solution")
add_body(
    "Idideration orchestrates a team of 8 specialized AI marketing agents \u2014 from a Behavioral "
    "Scientist who cites academic research on dopaminergic reward circuits, to a Psychometrics "
    "Expert who builds OCEAN personality profiles for each audience segment, to a Chief Strategist "
    "who synthesizes everything into a coherent grand strategy. Together, they produce a comprehensive, "
    "12-section marketing plan in hours that would take a traditional agency weeks and cost tens of "
    "thousands of dollars. Every recommendation is backed by data, citations, and cutting-edge "
    "behavioral science."
)

# The Agent Team
add_section_header("The Agent Team")

agents = [
    ("Intake Analyst", "Product assessment, theme extraction, market segments", "Flash"),
    ("Behavioral Scientist", "Neural pathways, behavioral loops, academic citations", "Pro"),
    ("Psychometrics Expert", "OCEAN profiles, persona cards, messaging DNA", "Pro"),
    ("Competitive Intelligence", "Competitor analysis, positioning, unclaimed territory", "Flash"),
    ("Social Strategist", "Platform strategy, content pillars, influencer mapping", "Flash"),
    ("Chief Strategist", "Grand strategy, multi-channel plan, measurement framework", "Pro"),
    ("Creative Director", "Creative brief, messaging architecture, CTAs, deliverables", "Pro"),
    ("Stakeholder Agent", "Interview question generation and answer collection", "Flash"),
]

table = doc.add_table(rows=len(agents) + 1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.6)
    row.cells[1].width = Inches(4.0)
    row.cells[2].width = Inches(0.6)

# Header row
headers = ["Agent", "Role", "Model"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A1A2E")
    set_cell_text(cell, h, size=Pt(9), bold=True, color=WHITE)

# Data rows
for idx, (agent, role, model) in enumerate(agents):
    row = table.rows[idx + 1]
    for ci, val in enumerate([agent, role, model]):
        cell = row.cells[ci]
        if idx % 2 == 1:
            set_cell_shading(cell, LIGHT_GRAY_BG)
        set_cell_text(cell, val, size=Pt(9))

# Remove table borders and add light ones
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr %s/>' % nsdecls('w'))
borders = parse_xml(
    '<w:tblBorders %s>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>' % nsdecls('w')
)
tblPr.append(borders)

# The Output
add_section_header("The Output")
add_body(
    "A 12-section DOCX marketing plan: Executive Summary, Product Assessment, Behavioral Framework, "
    "Audience Segmentation, Competitive Landscape, Social Strategy, Grand Strategy, Multi-Channel Plan, "
    "Creative Brief, Campaign Deliverables, Measurement Framework, and Research Citations. Every claim "
    "grounded in behavioral science with academic backing."
)

# ============================================================
# PAGE BREAK
# ============================================================
doc.add_page_break()

# ============================================================
# PAGE 2
# ============================================================

# Dual Knowledge Base
add_section_header("Dual Knowledge Base")
add_body("Two AI-accessible knowledge bases power every analysis:", space_after=Pt(2))

add_body_with_bold_label(
    "Marketing Bible (Global) \u2014 ",
    "40+ entries on psychometrics (Big Five/OCEAN, VALS), neuroscience (dopaminergic reward, "
    "prospect theory, mirror neurons), behavioral economics (loss aversion, anchoring, choice "
    "architecture), game theory (Nash equilibrium, network effects), social media strategy, "
    "content strategy, and brand building frameworks.",
    space_after=Pt(2)
)

add_body_with_bold_label(
    "Product Bible (Per-Project) \u2014 ",
    "All data specific to the product being marketed \u2014 notes, stakeholder input, competitive data, "
    "and for media products, scene-level content intelligence from Narralytica.",
    space_after=Pt(2)
)

# What Makes It Different
add_section_header("What Makes It Different")

diffs = [
    ("1. Behavioral Science Depth", " \u2014 Not just \u201cAI marketing copy.\u201d Every recommendation is grounded in neuroscience, psychometrics, and behavioral economics. The Behavioral Scientist agent cites academic papers and maps neural pathways."),
    ("2. Multi-Agent Collaboration", " \u2014 8 specialized agents build on each other\u2019s work, creating compound intelligence. The Chief Strategist synthesizes inputs from 6 prior agents \u2014 something no single human or generic AI can replicate."),
    ("3. $3 vs $30,000", " \u2014 A full 7-agent analysis costs approximately $3 in API fees. The equivalent agency deliverable costs $20-50K and takes weeks. This is a 10,000x cost reduction with greater scientific rigor."),
    ("4. Narralytica Integration", " \u2014 For media products, Idideration connects to scene-level video intelligence: 35+ metadata fields per scene, semantic search, mood/tone analytics. No other marketing tool understands content at this depth."),
]

for bold_part, normal_part in diffs:
    add_body_with_bold_label(bold_part, normal_part, space_after=Pt(2))

# First Case Study
add_section_header("First Case Study")
add_body(
    "Idideration\u2019s first deployment targets Homestead: The Series on Angel Studios \u2014 a genre-defining "
    "post-apocalyptic faith-based drama with a $20M+ theatrical run, S2 renewal, and 450K+ Guild "
    "community. The platform will deliver a full marketing plan including behavioral framework, OCEAN "
    "audience profiles, competitive positioning, and neuroscience-backed creative brief."
)

# Cost Structure
add_section_header("Cost Structure")

cost_data = [
    ("Per full analysis", "~$3 (Gemini API costs)"),
    ("Infrastructure", "Negligible (Hetzner VPS)"),
    ("Equivalent agency cost", "$20,000\u2013$50,000+"),
    ("ROI multiplier", "~10,000x"),
]

cost_table = doc.add_table(rows=len(cost_data), cols=2)
cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cost_table.autofit = True

for row in cost_table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(2.5)

for idx, (label, value) in enumerate(cost_data):
    row = cost_table.rows[idx]
    set_cell_text(row.cells[0], label, size=Pt(9), bold=True)
    set_cell_text(row.cells[1], value, size=Pt(9))
    if idx % 2 == 1:
        set_cell_shading(row.cells[0], LIGHT_GRAY_BG)
        set_cell_shading(row.cells[1], LIGHT_GRAY_BG)

# Light borders for cost table
cost_tbl = cost_table._tbl
cost_tblPr = cost_tbl.tblPr if cost_tbl.tblPr is not None else parse_xml('<w:tblPr %s/>' % nsdecls('w'))
cost_borders = parse_xml(
    '<w:tblBorders %s>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>' % nsdecls('w')
)
cost_tblPr.append(cost_borders)

# Built On Proven Architecture
add_section_header("Built On Proven Architecture")
add_body(
    "Idideration shares its agent orchestration architecture with Cassian (AI book editing) and "
    "integrates with Narralytica (video intelligence) \u2014 both built and operated by the same team."
)

add_hr()

# Footer
add_centered("idideration.com | AI-Powered Marketing Strategy", 8, color=GRAY_888, space_after=Pt(0))

# Save
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)

size = os.path.getsize(OUTPUT)
print(f"Created: {OUTPUT}")
print(f"Size: {size:,} bytes ({size/1024:.1f} KB)")
