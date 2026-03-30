#!/usr/bin/env python3
"""Generate the Idideration Product Brief DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = "/Users/JERS/Idideration/projects/Idideration_Product_Brief.docx"
DARK_BLUE = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
TABLE_HEADER_BG = "1A1A2E"
TABLE_ALT_BG = "F2F2F2"

doc = Document()

# ── Global styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = DARK_BLUE
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(24)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# 1-inch margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ── Helper functions ───────────────────────────────────────────────
def add_page_break():
    doc.add_page_break()

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.add_run(text).font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p

def add_body_bold(bold_part, normal_part):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
    return table


# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IDIDERATION")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Marketing Strategy Platform")
run.font.size = Pt(18)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Replacing Marketing Departments with Intelligence")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Product Brief \u2014 March 2026")
run.font.size = Pt(14)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Confidential")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS placeholder
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Vision & Mission",
    "2. The Problem",
    "3. The Solution",
    "4. The Agent Team",
    "5. Technology Architecture",
    "6. Knowledge Base System",
    "7. Workflow",
    "8. The Output \u2014 Marketing Plan",
    "9. Narralytica Integration",
    "10. Competitive Landscape",
    "11. Business Model",
    "12. Product Types Supported",
    "13. Cost Analysis",
    "14. Roadmap",
    "15. First Case Study \u2014 Homestead: The Series",
    "Appendix A: Glossary",
    "Appendix B: Technical Specifications",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 1: VISION & MISSION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Vision & Mission", level=1)

doc.add_heading("Vision", level=2)
add_body("Replace entire marketing departments with CMO-down capabilities using AI as a force multiplier. Idideration is not a tool that assists marketers\u2014it is the marketing department. From strategic planning to audience segmentation to creative briefing, the platform delivers the full spectrum of marketing intelligence that traditionally requires a team of 10-20 specialists.")

doc.add_heading("Mission", level=2)
add_body("Deliver genuinely useful, comprehensive marketing strategy to creatives, studios, startups, and brands\u2014faster, cheaper, and more rigorous than traditional agencies. Idideration exists to democratize world-class marketing strategy, making it accessible to anyone with a product worth promoting, regardless of their budget or connections.")

doc.add_heading("Core Thesis", level=2)
add_body("The best marketing sits at the intersection of behavioral science, neuroscience, psychometrics, and creative storytelling. AI can synthesize these disciplines at scale in ways no human team can. While a traditional marketing team might have one person who studied psychology and another who understands social media algorithms, no single human can hold deep expertise across all of these fields simultaneously. A properly orchestrated team of AI agents can\u2014and does.")

doc.add_heading("Target Market", level=2)
add_body("First target market: Entertainment (TV shows, films, streaming content)\u2014expanding to CPGs, apps, brands, and ideas. The entertainment vertical was chosen because it benefits most from the Narralytica integration, providing scene-level content intelligence that no other marketing platform can access.")

add_body_bold("First client: ", "Homestead: The Series (Angel Studios)\u2014a post-apocalyptic faith-based drama with a $20M+ theatrical run, Season 2 renewed, and a 450,000+ member Guild community.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 2: THE PROBLEM
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. The Problem", level=1)

add_body("The marketing industry is broken in fundamental ways that affect every company, creator, and brand trying to reach an audience. The gap between what marketing could be\u2014if it leveraged the full depth of behavioral science, neuroscience, and psychometrics\u2014and what it actually delivers is enormous.")

add_bullet("Traditional marketing agencies are expensive ($15K\u2013$50K/month retainers), slow (weeks to months for strategy), and often superficial. Most deliverables are recycled templates with a fresh coat of paint, not genuine strategic thinking tailored to the product.")
add_bullet("Most marketing strategies lack scientific rigor\u2014based on gut feeling, not behavioral science. A CMO might \"feel\" that a campaign will resonate, but they rarely cite dopaminergic reward circuits, prospect theory, or OCEAN personality profiles to justify their approach.")
add_bullet("Small studios, indie creators, and startups can\u2019t afford CMO-level strategic thinking. The people who need great marketing the most are the ones least able to pay for it. A filmmaker who spent years making a great show shouldn\u2019t have to settle for amateur marketing because they can\u2019t afford a $30K/month agency.")
add_bullet("Marketing departments are siloed\u2014the social media person doesn\u2019t think like a neuroscientist, the strategist doesn\u2019t understand psychometrics, and the creative director has never read a paper on loss aversion. Each specialist optimizes their lane without understanding the whole.")
add_bullet("The gap between academic research (what actually drives human behavior) and marketing practice is massive. Decades of peer-reviewed research in behavioral economics, cognitive psychology, and social neuroscience exist\u2014yet most marketing plans never reference a single study.")
add_bullet("No existing tool synthesizes behavioral science, neuroscience, game theory, and psychometrics into actionable marketing plans. Tools like HubSpot help you execute. Jasper helps you write copy. But nothing helps you think at the level of a CMO who also happens to be a behavioral scientist.")

add_body("The result: most products\u2014even great ones\u2014are marketed with strategies that scratch the surface of human motivation. Audiences are treated as demographics rather than psychographic profiles. Campaigns are launched based on what competitors are doing rather than what behavioral science says will actually move people to action.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 3: THE SOLUTION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. The Solution", level=1)

add_body("Idideration is a web application that orchestrates a team of 8 specialized AI agents, each an expert in a specific marketing discipline. Together they produce a comprehensive marketing plan that would take a full marketing department weeks to create. The platform runs the equivalent of a full-day strategy offsite\u2014with behavioral scientists, psychometricians, competitive analysts, and creative directors all in the room\u2014in under 30 minutes for approximately $3.")

add_body("What makes Idideration fundamentally different from other AI marketing tools is that it doesn\u2019t just generate content. It thinks strategically. Each agent reads from a shared knowledge base of behavioral science, neuroscience, and psychometrics\u2014the Marketing Bible\u2014and applies that knowledge to the specific product at hand. The output isn\u2019t generic marketing advice; it\u2019s a deeply researched, scientifically grounded strategy customized to the product, its audience, and its competitive landscape.")

doc.add_heading("Dual Knowledge Base Architecture", level=2)

add_body_bold("Marketing Bible (Global): ", "A comprehensive knowledge base containing 40+ entries covering psychometrics (Big Five/OCEAN, VALS), neuroscience (dopaminergic reward, prospect theory, mirror neurons), behavioral economics (loss aversion, anchoring, choice architecture), game theory (Nash equilibrium, signaling, network effects), social media strategy, content strategy, and brand building. This is the shared intelligence that every agent draws from\u2014the accumulated wisdom of decades of behavioral research, distilled into structured entries that AI agents can reason over.")

add_body_bold("Product Bible (Per-Project): ", "All data specific to the product being marketed\u2014unstructured notes, stakeholder input, research findings, competitive data, and Narralytica content intelligence. The Product Bible grows as agents work: the Intake Analyst writes initial assessments, the Behavioral Scientist adds research citations, stakeholder answers feed in, and each subsequent agent builds on the cumulative knowledge. This ensures that later agents in the pipeline have richer context than earlier ones.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 4: THE AGENT TEAM
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. The Agent Team", level=1)

add_body("Idideration\u2019s core innovation is its multi-agent architecture. Rather than using a single AI model to produce a marketing plan, the platform orchestrates 8 specialized agents in a deliberate sequence, each building on the work of those before it. This mirrors how the best marketing departments operate\u2014with specialists who hand off to each other\u2014while eliminating the communication failures and ego conflicts that plague human teams.")

# Agent table
agent_headers = ["Agent Name", "Expertise", "Model", "Key Outputs"]
agent_rows = [
    ["Intake Analyst", "Product assessment & initial analysis", "Gemini 2.5 Flash", "Product summary, core themes, emotional hooks, market segments, quick wins"],
    ["Behavioral Scientist", "Neuroscience & psychology", "Gemini 2.5 Pro", "Literature review with citations, behavioral framework, neural pathway mapping, motivational drivers, behavioral loops, framing recommendations"],
    ["Psychometrics Expert", "Audience modeling & segmentation", "Gemini 2.5 Pro", "OCEAN personality profiles per segment, VALS mapping, Schwartz values, persona cards with narratives, messaging DNA, prioritization matrix"],
    ["Competitive Intelligence", "Market landscape analysis", "Gemini 2.5 Flash", "Competitor profiles, synthesis takeaways, positioning recommendations, unclaimed territory identification"],
    ["Social Strategist", "Platform & community strategy", "Gemini 2.5 Flash", "Platform-specific content strategies, influencer mapping (macro/micro/nano), aggregator communities, hashtag analysis"],
    ["Chief Strategist", "Strategic synthesis (the AI CMO)", "Gemini 2.5 Pro", "Grand strategy (big idea, thesis, pillars, win conditions), timeline phases, multi-channel plan with budget allocation, measurement framework"],
    ["Creative Director", "Creative execution", "Gemini 2.5 Pro", "Creative brief, messaging architecture per segment with samples, CTA library by funnel stage, tagline options, campaign deliverables spec, do\u2019s and don\u2019ts"],
    ["Stakeholder Agent", "Interview management", "Gemini 2.5 Flash", "Targeted questions organized by role/purpose, answer collection workflow"],
]
make_table(agent_headers, agent_rows)

doc.add_paragraph()  # spacing

# ── Agent deep dives ──
doc.add_heading("Intake Analyst", level=2)
add_body("The Intake Analyst is the first agent to engage with any new project and serves as the foundation layer for everything that follows. It ingests all raw product data\u2014descriptions, notes, URLs, Narralytica content intelligence\u2014and produces a structured product summary that standardizes the information for downstream agents. Its outputs include core thematic analysis (identifying the 3-5 central themes of the product), emotional hook identification (what will make people feel something), initial market segment hypotheses, and quick-win opportunities that could generate immediate traction.")
add_body("What makes the Intake Analyst\u2019s output unique is its emphasis on emotional resonance and thematic depth rather than surface-level product features. While a human marketing coordinator might summarize a TV show as \u201cpost-apocalyptic drama,\u201d the Intake Analyst identifies specific emotional hooks like \u201cfamily bonds under existential threat\u201d and \u201cthe tension between self-preservation and community responsibility.\u201d These nuanced themes become the raw material that the Behavioral Scientist and Psychometrics Expert use to build their frameworks.")
add_body("When a Narralytica link is provided for media products, the Intake Analyst automatically fetches scene-level data\u201435+ metadata fields per scene\u2014giving it an unprecedented level of content understanding. It can identify which characters drive the most emotional resonance, which scenes contain the highest-impact moments for marketing, and which cultural themes run through the content.")

doc.add_heading("Behavioral Scientist", level=2)
add_body("The Behavioral Scientist is arguably the most distinctive agent in the Idideration system, and the one that most clearly differentiates the platform from anything else on the market. This agent reads the Intake Analyst\u2019s product summary, then conducts a systematic literature review against the Marketing Bible\u2019s neuroscience and behavioral economics entries. It maps the product\u2019s themes and emotional hooks to specific neural pathways, cognitive biases, and motivational frameworks.")
add_body("Key outputs include a behavioral framework with motivational drivers mapped to specific research (e.g., \u201cthe show\u2019s family survival themes activate oxytocin-mediated trust and in-group bonding circuits\u201d), neural pathway mapping that connects product attributes to brain regions and neurotransmitter systems, behavioral loop design (trigger \u2192 action \u2192 variable reward \u2192 investment) that shows how to create habit-forming engagement, and framing recommendations grounded in prospect theory and loss aversion research.")
add_body("Every claim the Behavioral Scientist makes is backed by a research citation with an evidence strength rating. This isn\u2019t pseudo-science\u2014it\u2019s the application of peer-reviewed behavioral research to marketing strategy. The agent uses Gemini 2.5 Pro because it needs deep reasoning capabilities to synthesize complex scientific concepts and apply them to a specific product context.")

doc.add_heading("Psychometrics Expert", level=2)
add_body("The Psychometrics Expert takes the behavioral framework and translates it into concrete audience profiles. Using the Big Five/OCEAN personality model, VALS psychographic segmentation, Schwartz\u2019s Theory of Basic Human Values, and Moral Foundations Theory, this agent creates detailed persona cards for each target segment\u2014complete with \u201cday in the life\u201d narratives that bring the data to life.")
add_body("The agent\u2019s most powerful output is the Messaging DNA matrix: for each audience segment, it specifies the optimal tone, complexity level, activating language, reactance triggers to avoid, and emotional valence. This goes far beyond traditional demographics. Instead of \u201cwomen 25-45,\u201d you get \u201chigh-Openness, high-Agreeableness individuals who score strongly on Universalism and Benevolence values, respond to narrative-driven messaging with moderate complexity, and are triggered to reactance by heavy-handed sales language.\u201d")
add_body("The prioritization matrix ranks segments by accessibility (how easy to reach), receptivity (how likely to respond), advocacy potential (how likely to share), and lifetime value. This tells the Chief Strategist where to allocate resources for maximum impact. The Psychometrics Expert uses Gemini 2.5 Pro because persona construction and psychometric modeling require sophisticated reasoning over multiple overlapping frameworks.")

doc.add_heading("Competitive Intelligence", level=2)
add_body("The Competitive Intelligence agent maps the product\u2019s competitive landscape by analyzing direct competitors, adjacent alternatives, and substitute behaviors. For each competitor, it produces a structured profile covering brand style, target audiences, social media presence, notable campaigns, and estimated financial performance. The agent draws from the Marketing Bible\u2019s game theory entries\u2014particularly signaling theory and positioning strategy\u2014to identify how competitors are positioning themselves and where gaps exist.")
add_body("The most valuable output is the unclaimed territory analysis: specific positioning spaces, audience segments, or messaging approaches that no competitor currently occupies. This directly informs the Chief Strategist\u2019s grand strategy by revealing where the product can establish a unique, defensible position rather than fighting for crowded ground.")
add_body("The agent uses Gemini 2.5 Flash because competitive analysis is primarily a data-gathering and pattern-recognition task rather than a deep-reasoning task. Flash\u2019s speed allows it to process multiple competitors efficiently while keeping costs low.")

doc.add_heading("Social Strategist", level=2)
add_body("The Social Strategist translates the behavioral framework and audience profiles into platform-specific content strategies. For each major platform (Instagram, TikTok, YouTube, Reddit, X, Facebook), it defines content pillars, recommended formats, posting frequency, engagement tactics, and growth strategies\u2014all calibrated to the platform\u2019s current algorithm preferences and the target audience\u2019s platform behavior.")
add_body("The agent\u2019s influencer mapping capability segments potential partners into macro (500K+), micro (10K-500K), and nano (1K-10K) tiers, with specific criteria for each tier based on the product\u2019s audience profiles. It also identifies aggregator communities\u2014subreddits, Facebook groups, Discord servers, forums\u2014where the target audience already gathers, providing a roadmap for organic community engagement.")
add_body("The Social Strategist draws heavily from the Marketing Bible\u2019s social media entries, which are updated to reflect current (2025) platform algorithms and best practices. It uses Gemini 2.5 Flash because social media strategy requires breadth of knowledge across many platforms rather than the deep reasoning needed for behavioral analysis.")

doc.add_heading("Chief Strategist", level=2)
add_body("The Chief Strategist is the AI CMO\u2014the agent that synthesizes all preceding work into a unified grand strategy. It reads the outputs of every agent that ran before it and produces the strategic backbone of the entire marketing plan: the big idea, strategic thesis, pillar framework, win conditions, timeline phases, multi-channel plan with budget allocation, and measurement framework.")
add_body("What makes the Chief Strategist\u2019s output exceptional is its integration depth. It doesn\u2019t just list channels and tactics; it explains why each element exists, grounding every recommendation in the behavioral framework, audience profiles, and competitive positioning. A typical output might read: \u201cAllocate 40% of Phase 1 budget to short-form video (TikTok/Reels) because Segment A\u2019s high-Openness, high-Sensation-Seeking profile correlates with 3.2x higher engagement on visual-first platforms, and the competitive landscape shows no competitor effectively occupying the faith-meets-survival content niche on TikTok.\u201d")
add_body("The measurement framework includes a north star metric, leading indicators tied to each strategic pillar, and a reporting cadence. This ensures that the strategy isn\u2019t just a document that sits on a shelf\u2014it\u2019s an actionable plan with clear success criteria. The Chief Strategist uses Gemini 2.5 Pro because strategic synthesis is the most reasoning-intensive task in the entire pipeline.")

doc.add_heading("Creative Director", level=2)
add_body("The Creative Director takes the grand strategy and translates it into creative execution. Its outputs include a comprehensive creative brief (target audience, key message, tone, mandatories, deliverables), messaging architecture for each audience segment with sample copy, a CTA library organized by funnel stage (awareness, consideration, conversion, retention), tagline options, campaign deliverables specification, and a detailed do\u2019s and don\u2019ts guide.")
add_body("The messaging architecture is particularly powerful: for each audience segment identified by the Psychometrics Expert, the Creative Director produces sample headlines, body copy, and CTAs calibrated to that segment\u2019s messaging DNA. The difference between \u201cWatch now\u201d and \u201cJoin 450,000 families who\u2019ve already discovered Homestead\u201d is the difference between generic and psychometrically informed creative\u2014and the Creative Director consistently produces the latter.")
add_body("The agent uses Gemini 2.5 Pro because creative execution at this level requires sophisticated reasoning about audience psychology, brand voice, and strategic alignment. It\u2019s not generating random copy\u2014it\u2019s producing creative work that is directly traceable to behavioral research and audience profiles.")

doc.add_heading("Stakeholder Agent", level=2)
add_body("The Stakeholder Agent serves a unique role in the pipeline: rather than producing analysis, it produces questions. Based on the product data and (optionally) the outputs of other agents, it generates targeted interview questions organized by stakeholder role (creator, producer, marketing lead, talent) and purpose (brand positioning, audience insight, competitive intelligence, content strategy).")
add_body("The questions are designed to extract the specific information that would most improve the other agents\u2019 outputs. For example, if the Intake Analyst identified a theme of \u201cfamily resilience\u201d but lacked data on the creator\u2019s intent, the Stakeholder Agent would generate questions like: \u201cWhat specific real-world family dynamics inspired the Bennett family\u2019s storyline?\u201d and \u201cHow do you want audiences to feel about their own families after watching?\u201d")
add_body("Stakeholder answers feed directly back into the Product Bible, enriching the context for subsequent agent runs. This creates a feedback loop: run agents \u2192 identify knowledge gaps \u2192 interview stakeholders \u2192 enrich Product Bible \u2192 re-run agents with deeper context. The Stakeholder Agent uses Gemini 2.5 Flash because question generation is a structured task that benefits from speed rather than deep reasoning.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 5: TECHNOLOGY ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Technology Architecture", level=1)

add_body("Idideration is built on a modern Python web stack optimized for AI orchestration, real-time monitoring, and professional document generation. The architecture mirrors Cassian (a book editing AI platform by the same creator), a proven pattern for multi-agent AI applications.")

doc.add_heading("Core Stack", level=2)
add_bullet("Backend: FastAPI (Python) with SQLAlchemy ORM", bold_prefix="")
add_bullet("Database: SQLite (local development), PostgreSQL (production)", bold_prefix="")
add_bullet("AI Engine: Google Gemini 2.5 (Flash for data-gathering agents, Pro for reasoning-heavy agents)", bold_prefix="")
add_bullet("Frontend: Server-rendered Jinja2 templates with HTMX for dynamic updates", bold_prefix="")
add_bullet("Knowledge Base: Dual-bible system (Marketing Bible global + Product Bible per-project)", bold_prefix="")
add_bullet("Output: python-docx for comprehensive DOCX marketing plans", bold_prefix="")
add_bullet("Integration: Narralytica API for scene-level video intelligence on media products", bold_prefix="")
add_bullet("Deployment: Hetzner VPS, nginx reverse proxy, systemd service, Let\u2019s Encrypt HTTPS", bold_prefix="")
add_bullet("Domain: idideration.com", bold_prefix="")

doc.add_heading("Why This Stack", level=2)
add_body("FastAPI was chosen for its async support (critical for long-running AI agent pipelines), automatic API documentation, and Python ecosystem compatibility. HTMX provides reactive UI updates without the complexity of a JavaScript framework\u2014important because the primary interface complexity is in monitoring agent progress, not in rich client-side interactions. Jinja2 templates keep the frontend simple and server-rendered, reducing deployment complexity.")
add_body("Google Gemini 2.5 was selected as the AI engine for its combination of large context windows (essential for agents that need to read entire knowledge bases), strong reasoning capabilities in the Pro tier, cost-effective data gathering in the Flash tier, and competitive pricing that keeps per-run costs under $5. The dual-model approach\u2014Flash for data-gathering agents, Pro for reasoning-heavy agents\u2014optimizes the cost-quality tradeoff across the pipeline.")

doc.add_heading("Data Model", level=2)
add_body("The data model is organized around these key entities:")
add_bullet("Users \u2014 Authentication, permissions, and project ownership", bold_prefix="")
add_bullet("Projects \u2014 Each product being marketed, with type, description, raw data, and optional Narralytica link", bold_prefix="")
add_bullet("Bible Entries (18 categories) \u2014 Both Marketing Bible (global) and Product Bible (per-project) entries, categorized and tagged for agent retrieval", bold_prefix="")
add_bullet("Crew Runs \u2014 A single execution of the full agent pipeline, with status tracking, timestamps, and aggregate cost", bold_prefix="")
add_bullet("Agent Runs \u2014 Individual agent executions within a crew run, with input/output token counts, cost tracking, timing, and generated output", bold_prefix="")
add_bullet("Stakeholder Questions \u2014 Generated interview questions with role targeting and answer collection status", bold_prefix="")
add_bullet("Research Citations \u2014 Academic papers and sources referenced by agents, with evidence strength ratings", bold_prefix="")
add_bullet("Output Documents \u2014 Generated DOCX marketing plans linked to their source crew run", bold_prefix="")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 6: KNOWLEDGE BASE SYSTEM
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. Knowledge Base System", level=1)

add_body("The knowledge base system is Idideration\u2019s intellectual foundation\u2014the accumulated expertise that transforms generic AI outputs into genuinely insightful marketing strategy. It operates on a dual-bible architecture that separates universal marketing knowledge from project-specific intelligence.")

doc.add_heading("The Marketing Bible", level=2)
add_body("The Marketing Bible contains 40 seed entries across 8 categories, representing the core disciplines that inform world-class marketing strategy. Each entry contains 2-4 substantive paragraphs plus structured data optimized for agent reference. Total seed content: approximately 125KB of dense, actionable knowledge.")

kb_headers = ["Category", "Entries", "Key Concepts"]
kb_rows = [
    ["Frameworks", "6", "AIDA, Hook Model, Jobs-to-Be-Done, Blue Ocean Strategy, Crossing the Chasm, STP"],
    ["Psychometrics", "5", "Big Five/OCEAN, VALS 8 segments, Schwartz 10 values, Moral Foundations, Regulatory Focus Theory"],
    ["Neuroscience", "6", "Dopaminergic reward circuits, Prospect Theory, Mirror neurons/narrative transportation, Oxytocin/trust, Cortisol/FOMO, Default mode network"],
    ["Behavioral Economics", "6", "Loss aversion/framing, Anchoring, Choice architecture/Nudge, Endowment/IKEA effect, Social proof/cascades, Hyperbolic discounting"],
    ["Social Media", "5", "Instagram 2025, TikTok 2025, YouTube algorithm, Reddit engagement, Cross-platform repurposing"],
    ["Game Theory", "4", "Nash equilibrium, Signaling theory, Network effects/viral coefficients, Mechanism design"],
    ["Content Strategy", "4", "Content pillars, Hero/Hub/Help, Content-market fit, Storytelling arc/Freytag"],
    ["Brand Building", "4", "Positioning statement template, Jung\u2019s 12 archetypes, Keller\u2019s CBBE pyramid, Byron Sharp distinctive assets"],
]
make_table(kb_headers, kb_rows)

doc.add_paragraph()

add_body("Each Marketing Bible entry is written to be both human-readable and agent-optimized. Entries include narrative explanations of the concept, structured data (e.g., the OCEAN model entry contains all five dimensions with their sub-facets and marketing implications), practical application guidelines, and cross-references to related entries. This structure allows agents to both understand the concept deeply and extract specific data points for their analysis.")

doc.add_heading("The Product Bible", level=2)
add_body("The Product Bible is a per-project knowledge base that accumulates intelligence about the specific product being marketed. Unlike the Marketing Bible, which is curated and static, the Product Bible is dynamic\u2014it grows throughout the agent pipeline as each agent writes its findings back to the project\u2019s knowledge base.")
add_body("Sources of Product Bible content include: manual entries added by the user (product descriptions, brand guidelines, audience research), imported data from Narralytica (scene-level content intelligence for media products), agent-generated entries (each agent can write structured findings to the Product Bible for downstream agents to access), and stakeholder interview answers (collected by the Stakeholder Agent and fed back into the bible).")
add_body("This cumulative architecture means that the last agent in the pipeline\u2014the Chief Strategist or Creative Director\u2014has access to a far richer context than the first agent. The Product Bible serves as the team\u2019s shared memory, ensuring that insights discovered by one agent are available to all subsequent agents.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 7: WORKFLOW
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. Workflow", level=1)

add_body("Idideration\u2019s workflow is designed to be simple for the user while orchestrating complex agent interactions behind the scenes. A typical project moves through the following steps:")

steps = [
    ("1. Create New Project: ", "Define the product with a name, type (TV series, film, CPG, app, brand, idea), description, raw data dump, and optionally link a Narralytica show ID for media products."),
    ("2. Populate Product Bible: ", "Add manual entries with product-specific knowledge\u2014brand guidelines, audience research, stakeholder notes, competitive observations. For media products with a Narralytica link, content intelligence is automatically imported."),
    ("3. Run Full Analysis: ", "Launch all 7 analysis agents in sequence (Intake \u2192 Behavioral Scientist \u2192 Psychometrics \u2192 Competitive Intelligence \u2192 Social Strategist \u2192 Chief Strategist \u2192 Creative Director), or select specific agents for targeted analysis."),
    ("4. Monitor Real-Time Progress: ", "Watch the agent pipeline execute through the web UI, with status indicators (queued, running, complete, failed) for each agent, live token counts, and running cost tallies."),
    ("5. Review Agent Outputs: ", "Examine each agent\u2019s output individually in the web interface. Outputs are rendered with full formatting, tables, and citations."),
    ("6. Generate DOCX Marketing Plan: ", "Compile all completed agent outputs into a comprehensive, professionally formatted DOCX marketing plan with 12 sections, tables, and research citations."),
    ("7. Download and Deliver: ", "Download the generated DOCX and deliver it to the client. The document is ready for presentation without additional formatting."),
    ("8. Run Stakeholder Agent (Optional): ", "Generate targeted interview questions based on the product data and analysis gaps. Questions are organized by stakeholder role and information purpose."),
    ("9. Feed Stakeholder Answers: ", "Collect stakeholder responses and feed them back into the Product Bible, enriching the project\u2019s knowledge base with insider context and creator intent."),
    ("10. Re-Run for Refined Strategy: ", "Re-run agents with the enriched Product Bible context. The cumulative knowledge produces deeper, more nuanced outputs on each successive run."),
]
for bold, text in steps:
    add_body_bold(bold, text)

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 8: THE OUTPUT
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. The Output \u2014 Marketing Plan", level=1)

add_body("The generated DOCX marketing plan is the primary deliverable of the Idideration platform. It is a comprehensive, professionally formatted document containing 12 sections that together constitute a complete marketing strategy. The document is designed to be client-ready\u2014suitable for presentation to stakeholders, executives, or creative teams without additional formatting or editing.")

output_sections = [
    ("1. Executive Summary", "Presents the big idea, strategic thesis, and win conditions in a concise overview. This section gives stakeholders a clear, compelling picture of the entire strategy in 2-3 pages, including the core insight that drives the plan and the measurable outcomes that define success."),
    ("2. Product Assessment", "Covers core themes, emotional hooks, initial market segments, and quick-win opportunities identified by the Intake Analyst. This section ensures all stakeholders share a common understanding of what the product is, what makes it special, and where the immediate opportunities lie."),
    ("3. Behavioral Framework", "Details motivational drivers, neural pathway mapping, behavioral loops (trigger \u2192 action \u2192 variable reward \u2192 investment), and framing recommendations. This is the scientific backbone of the strategy\u2014the section that differentiates an Idideration plan from any traditional agency deliverable."),
    ("4. Audience Segmentation", "Contains OCEAN personality profiles for each target segment, messaging DNA specifications, persona cards with day-in-the-life narratives, and a prioritization matrix. This section transforms abstract audience data into vivid, actionable profiles that creative teams can design for."),
    ("5. Competitive Landscape", "Presents competitor profiles with brand style, audiences, social presence, and campaign analysis, plus positioning recommendations and unclaimed territory identification. This section shows where the product fits in the market and where it can carve out defensible space."),
    ("6. Social Media Strategy", "Delivers platform-specific content strategies with pillars, formats, and frequency; influencer mapping across macro, micro, and nano tiers; and aggregator community identification. Each platform strategy is calibrated to current algorithm preferences and audience behavior."),
    ("7. Grand Strategy", "Presents the strategic pillars, objectives with behavioral basis, phase-based timeline, and the overarching strategic thesis that unifies all tactical recommendations. This is the CMO-level strategic thinking that ties everything together."),
    ("8. Multi-Channel Plan", "Specifies channels, their roles in the strategy, budget allocation percentages, and KPIs for each channel. This section translates strategy into a practical execution plan with clear resource allocation."),
    ("9. Creative Brief", "Provides creative strategy, direction, and messaging architecture per audience segment with sample copy. This section gives creative teams everything they need to produce on-brand, psychometrically informed content."),
    ("10. Campaign Deliverables", "Lists tagline options, CTA library organized by funnel stage, and a comprehensive do\u2019s and don\u2019ts guide. This section ensures consistency across all campaign touchpoints and prevents common messaging mistakes."),
    ("11. Measurement Framework", "Defines the north star metric, leading indicators tied to each strategic pillar, and reporting cadence. This section ensures the strategy is measurable and accountable, not just a document that sits on a shelf."),
    ("12. Appendix: Research Citations", "Compiles all academic papers and research sources referenced throughout the plan, with evidence strength ratings. This section provides the scientific credibility that underpins the entire strategy."),
]
for title, desc in output_sections:
    add_body_bold(f"{title}: ", desc)

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 9: NARRALYTICA INTEGRATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("9. Narralytica Integration", level=1)

add_body("For media products (TV series, films, streaming content), Idideration connects to Narralytica\u2014a sister platform providing scene-level video intelligence that gives Idideration content understanding no other marketing tool can match. Narralytica is a separate video intelligence platform built by the same team, and its integration with Idideration creates a unique competitive advantage in the entertainment marketing space.")

doc.add_heading("Narralytica Capabilities", level=2)
add_bullet("35+ metadata fields per scene, including characters present, dialog summary, mood, tone, pacing, cultural references, content safety flags, and thematic tags")
add_bullet("Semantic search across episodes\u2014e.g., \u201cfind all scenes with emotional family moments\u201d or \u201cscenes featuring the protagonist in danger\u201d")
add_bullet("Mood timeline and emotional arc analytics, visualizing the emotional journey of an episode or season")
add_bullet("Character screen time data with scene-level granularity")
add_bullet("CutPrint scene detection technology for precise scene boundary identification")

doc.add_heading("How Integration Works", level=2)
add_body("When a project is created with a Narralytica show ID, the Intake Analyst automatically fetches scene-level data through the Narralytica API. This gives the agent access to the actual content of the show\u2014not just a synopsis or logline, but scene-by-scene understanding of characters, emotions, themes, and narrative structure.")
add_body("This content intelligence flows through the entire agent pipeline. The Behavioral Scientist can identify which specific scenes trigger dopaminergic reward circuits. The Psychometrics Expert can map character arcs to audience personality profiles. The Social Strategist can recommend specific scene clips for social media content. The Creative Director can reference actual moments from the show in campaign concepts.")
add_body("No other marketing platform in the world has access to this level of content intelligence. Traditional agencies watch a show and take notes. Idideration reads 35+ structured metadata fields for every scene, then reasons over that data with behavioral science, psychometrics, and competitive intelligence. This is the Narralytica advantage.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 10: COMPETITIVE LANDSCAPE (IDIDERATION)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Competitive Landscape", level=1)

add_body("Idideration occupies a unique position in the marketing technology landscape. No existing product combines multi-agent AI orchestration, behavioral science depth, psychometric audience modeling, and content intelligence integration. The following comparison maps Idideration against the major categories of marketing solutions.")

comp_headers = ["Competitor Type", "Examples", "Strengths", "Weaknesses", "Idideration Advantage"]
comp_rows = [
    ["Traditional Agencies", "WPP, Dentsu, boutique firms", "Human creativity, relationships", "$15-50K/mo, slow, no science", "10,000x cheaper, behavioral science depth"],
    ["AI Writing Tools", "Jasper, Copy.ai, ChatGPT", "Fast copy generation", "No strategy, no audience modeling", "Full CMO-level strategy, not just copy"],
    ["Marketing Platforms", "HubSpot, Hootsuite", "Execution & tracking", "Tools not strategy", "Decides what to do, not just helps do it"],
    ["Strategy Consultants", "McKinsey, Bain", "Deep analysis", "$500K+ engagements, months", "Same rigor, hours not months, fraction of cost"],
    ["Research Tools", "SEMrush, SimilarWeb", "Data collection", "No synthesis, no behavioral science", "Synthesizes data into actionable strategy"],
]
make_table(comp_headers, comp_rows)

doc.add_paragraph()

doc.add_heading("Idideration\u2019s Moat", level=2)
add_body("The platform\u2019s defensible advantage lies in the combination of four elements that no competitor replicates: (1) behavioral science depth embedded in the Marketing Bible and agent prompts, (2) multi-agent collaboration that mirrors a full marketing department, (3) the dual knowledge base system that accumulates project intelligence, and (4) Narralytica integration that provides content intelligence unavailable to any other marketing tool. Each element alone is valuable; together they create a product that operates in a category of one.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 11: BUSINESS MODEL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("11. Business Model", level=1)

add_body_bold("Phase 1 (Current): ", "Internal tool for serving clients directly. Idideration is used to produce marketing plans that are delivered to clients as a service. This allows the platform to be refined through real-world usage while generating revenue from day one.")
add_body_bold("Phase 2 (Future): ", "Potential SaaS for agencies, studios, and brands. Once the platform is proven through client work, it could be offered as a self-service tool, allowing agencies to run their own analyses and studios to generate marketing plans in-house.")

doc.add_heading("Revenue Model", level=2)
add_body("Per-project fees, dramatically undercutting traditional agency costs. A full marketing strategy that would cost $20,000\u2013$50,000+ from a traditional agency can be delivered for a fraction of that price while providing greater scientific rigor and depth.")

doc.add_heading("Unit Economics", level=2)
add_body("Cost per full agent run: approximately $3.10 (Gemini API costs). Equivalent agency cost for the same scope of work: $20,000\u2013$50,000+. This represents a roughly 10,000x cost reduction on the production side, enabling aggressive pricing that is simultaneously profitable for Idideration and transformatively affordable for clients.")

add_body("The value proposition is not just cost reduction\u2014it\u2019s the combination of cost, speed, and rigor. A $3 run that takes 30 minutes and produces strategy grounded in behavioral science is categorically different from a $30,000 engagement that takes 6 weeks and produces strategy grounded in gut feeling.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 12: PRODUCT TYPES SUPPORTED
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("12. Product Types Supported", level=1)

add_body("Idideration is designed to produce marketing strategy for any product that needs to find and reach an audience. The platform\u2019s agent architecture and knowledge base are product-type agnostic\u2014the same behavioral science, psychometrics, and strategic frameworks apply whether marketing a TV show, a consumer product, or a startup.")

types = [
    ("TV Series (first priority)", "The initial and primary use case, leveraging Narralytica integration for scene-level content intelligence. First test case: Homestead: The Series on Angel Studios."),
    ("Films", "Feature films benefit from the same content intelligence pipeline and behavioral framework, with adjustments for theatrical release windows and campaign timing."),
    ("Consumer Packaged Goods (CPGs)", "Physical products benefit from the psychometric audience modeling and behavioral economics frameworks, particularly loss aversion, choice architecture, and social proof cascades."),
    ("Apps & Software", "Digital products leverage the Hook Model behavioral loops, network effects analysis from game theory, and platform-specific social media strategies."),
    ("Brands & Companies", "Corporate brand strategy benefits from archetype mapping, positioning analysis, and the full competitive intelligence pipeline."),
    ("Ideas & Concepts", "Pre-product concepts\u2014campaigns, movements, causes\u2014can use Idideration to develop positioning, audience segmentation, and launch strategy before the product itself exists."),
]
for title, desc in types:
    add_body_bold(f"{title}: ", desc)

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 13: COST ANALYSIS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("13. Cost Analysis", level=1)

add_body("One of Idideration\u2019s most compelling attributes is its cost structure. The entire agent pipeline\u2014producing strategy that would take a marketing department weeks and cost tens of thousands of dollars\u2014runs for approximately $3.10 in API costs.")

cost_headers = ["Agent", "Model", "Est. Input Tokens", "Est. Output Tokens", "Est. Cost"]
cost_rows = [
    ["Intake Analyst", "Flash", "~50K", "~8K", "~$0.10"],
    ["Behavioral Scientist", "Pro", "~80K", "~15K", "~$0.80"],
    ["Psychometrics Expert", "Pro", "~90K", "~12K", "~$0.60"],
    ["Competitive Intelligence", "Flash", "~60K", "~10K", "~$0.10"],
    ["Social Strategist", "Flash", "~70K", "~10K", "~$0.10"],
    ["Chief Strategist", "Pro", "~120K", "~15K", "~$0.80"],
    ["Creative Director", "Pro", "~100K", "~12K", "~$0.60"],
    ["TOTAL", "\u2014", "~570K", "~82K", "~$3.10"],
]
make_table(cost_headers, cost_rows)

doc.add_paragraph()

add_body("To put this in perspective: $3.10 per run versus $20,000\u2013$50,000+ for equivalent agency work. Even accounting for platform overhead, staff time, and margin, Idideration can deliver comprehensive marketing strategy at a price point that makes world-class strategic thinking accessible to any creator, studio, or brand.")

add_body("The token estimates reflect full knowledge base loading (Marketing Bible + Product Bible), agent-specific prompt engineering, and comprehensive output generation. Actual costs vary by project complexity and knowledge base size, but the order of magnitude remains consistent: single-digit dollars for what traditionally costs five figures.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 14: ROADMAP
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("14. Roadmap", level=1)

doc.add_heading("Phase 1 \u2014 Current", level=2)
add_body("Core platform operational with all foundational capabilities in place:")
add_bullet("8 specialized AI agents fully implemented and tested")
add_bullet("Dual knowledge base system (Marketing Bible with 40 seed entries + Product Bible per-project)")
add_bullet("DOCX output generation with professional formatting and 12-section structure")
add_bullet("Narralytica API integration for scene-level content intelligence")
add_bullet("Real-time agent monitoring with status, token counts, and cost tracking")
add_bullet("Stakeholder question generation and answer collection workflow")
add_bullet("First case deployment: Homestead: The Series (Angel Studios)")

doc.add_heading("Phase 2 \u2014 Near-Term", level=2)
add_body("Expanding capabilities and efficiency:")
add_bullet("Web search integration for real-time competitive intelligence and trend data")
add_bullet("NotebookLM API integration for enhanced research synthesis")
add_bullet("Agent parallelization: Agents 2+3 (Behavioral Scientist + Psychometrics) and 4+5 (Competitive Intelligence + Social Strategist) running concurrently to reduce total pipeline time")
add_bullet("PDF output option alongside DOCX")
add_bullet("Slide deck generation for presentation-ready strategy decks")

doc.add_heading("Phase 3 \u2014 Future", level=2)
add_body("Scaling the platform:")
add_bullet("Multi-user support with role-based access control")
add_bullet("Client portal for stakeholder interview collection (self-service question answering)")
add_bullet("Cross-project analysis and pattern detection across the client portfolio")
add_bullet("Template library for common product types (TV series, film, CPG, app)")
add_bullet("External API for third-party integration and white-label usage")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 15: FIRST CASE STUDY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("15. First Case Study \u2014 Homestead: The Series", level=1)

add_body("Homestead: The Series serves as Idideration\u2019s first full deployment\u2014the proving ground where the platform\u2019s capabilities are tested against a real product with real stakes. The show provides an ideal test case because it has a passionate existing audience, a unique market position, and enough complexity to exercise every agent in the pipeline.")

doc.add_heading("About Homestead", level=2)
add_bullet("Post-apocalyptic faith-based drama distributed by Angel Studios")
add_bullet("$20M+ theatrical gross, Season 2 renewed")
add_bullet("450,000+ member Guild community (Angel Studios\u2019 crowdfunding/community platform)")
add_bullet("Based on the Black Autumn book series by Jeff Kirkham, an actual Green Beret")
add_bullet("Cast includes Bailey Chase, Charles Esten, Dawn Olivieri, and Jill Wagner")
add_bullet("Unique positioning: the first post-apocalyptic show created specifically for the faith audience")

doc.add_heading("Why Homestead Is the Ideal First Case", level=2)
add_body("Homestead occupies a fascinating market position: it\u2019s a genre show (post-apocalyptic survival) for an audience (faith-based viewers) that has historically been underserved by that genre. This creates rich opportunities for the Behavioral Scientist (survival instinct psychology meets faith-based motivational frameworks), the Psychometrics Expert (OCEAN profiles for faith audiences who also enjoy genre content), and the Competitive Intelligence agent (mapping both the faith entertainment space and the post-apocalyptic genre space).")
add_body("The Narralytica integration is particularly powerful for Homestead, providing scene-level intelligence about character dynamics, emotional arcs, family moments, faith themes, and survival tension that directly informs the marketing strategy. Idideration will produce a comprehensive marketing plan covering the full behavioral framework, OCEAN audience profiles, competitive positioning, multi-channel strategy, and creative brief\u2014demonstrating the platform\u2019s end-to-end capability on a real, high-stakes project.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX A: GLOSSARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Appendix A: Glossary", level=1)

glossary = [
    ("Marketing Bible", "The global knowledge base containing 40+ entries on behavioral science, neuroscience, psychometrics, game theory, social media strategy, content strategy, and brand building. Shared across all projects."),
    ("Product Bible", "A per-project knowledge base that accumulates intelligence about the specific product being marketed. Grows dynamically as agents write findings and stakeholder answers are collected."),
    ("Agent", "A specialized AI module configured with a specific expertise area, knowledge base access pattern, and output structure. Each agent uses either Gemini 2.5 Flash (data-gathering) or Gemini 2.5 Pro (deep reasoning)."),
    ("Crew Run", "A single execution of the full agent pipeline (or selected agents) for a project. Tracks aggregate status, timing, token usage, and cost."),
    ("OCEAN / Big Five", "The five-factor model of personality: Openness, Conscientiousness, Extraversion, Agreeableness, and Neuroticism. Used by the Psychometrics Expert to build audience profiles that predict messaging preferences and behavioral responses."),
    ("VALS", "Values and Lifestyles psychographic segmentation system. Classifies consumers into 8 segments (Innovators, Thinkers, Achievers, Experiencers, Believers, Strivers, Makers, Survivors) based on primary motivation and resources."),
    ("Behavioral Loop", "A cycle of trigger \u2192 action \u2192 variable reward \u2192 investment, based on Nir Eyal\u2019s Hook Model. Used by the Behavioral Scientist to design habit-forming engagement patterns."),
    ("Framing", "The presentation of information in a way that influences decision-making, grounded in Kahneman and Tversky\u2019s research. Loss framing (\u201cdon\u2019t miss out\u201d) vs. gain framing (\u201cdiscover something new\u201d) produces measurably different behavioral responses."),
    ("Prospect Theory", "Daniel Kahneman and Amos Tversky\u2019s theory that people evaluate outcomes relative to a reference point and weight losses more heavily than equivalent gains. Foundational to the Behavioral Scientist\u2019s framing recommendations."),
    ("Hook Model", "Nir Eyal\u2019s framework for building habit-forming products: Trigger (internal/external) \u2192 Action (simplest behavior) \u2192 Variable Reward (satisfaction) \u2192 Investment (stored value). Applied to marketing engagement design."),
    ("CTA (Call to Action)", "A specific prompt directing the audience to take a desired action. The Creative Director generates CTA libraries organized by funnel stage (awareness, consideration, conversion, retention)."),
    ("North Star Metric", "The single metric that best captures the core value the product delivers to customers. Defined by the Chief Strategist as the primary success indicator for the marketing strategy."),
    ("BISAC", "Book Industry Standards and Communications subject codes. Used for categorization of book-related products in the Product Bible."),
    ("Narralytica", "A sister platform providing scene-level video intelligence with 35+ metadata fields per scene, semantic search, mood timelines, and character analytics. Integrated with Idideration for media product marketing."),
    ("CutPrint", "Narralytica\u2019s proprietary scene detection technology that identifies scene boundaries in video content with high precision, enabling scene-level metadata extraction."),
]
for term, definition in glossary:
    add_body_bold(f"{term}: ", definition)

add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX B: TECHNICAL SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Appendix B: Technical Specifications", level=1)

doc.add_heading("Runtime & Frameworks", level=2)
add_bullet("Python 3.12+")
add_bullet("FastAPI web framework with Uvicorn ASGI server")
add_bullet("SQLAlchemy 2.0 ORM with async session support")
add_bullet("Google Gemini 2.5 (Flash and Pro tiers)")
add_bullet("Jinja2 templating engine")
add_bullet("HTMX for dynamic frontend updates")
add_bullet("python-docx for document generation")

doc.add_heading("Ports", level=2)
add_bullet("Production: Port 8011")
add_bullet("Local development: Port 8006")

doc.add_heading("Deployment", level=2)
add_bullet("Systemd service: idideration.service")
add_bullet("Nginx reverse proxy with SSL termination")
add_bullet("Let\u2019s Encrypt HTTPS certificates with auto-renewal")
add_bullet("Auto-restart on failure (RestartSec=5)")
add_bullet("Hetzner VPS hosting")
add_bullet("Domain: idideration.com")

doc.add_heading("Database", level=2)
add_bullet("SQLite for local development (zero-config)")
add_bullet("PostgreSQL for production deployment")
add_bullet("Alembic for database migrations")

doc.add_heading("Security", level=2)
add_bullet("API key authentication for Gemini and Narralytica services")
add_bullet("Environment-based configuration (no secrets in code)")
add_bullet("HTTPS-only in production")
add_bullet("Systemd sandboxing and privilege restriction")

# ── Footer note ────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u2014 End of Document \u2014")
run.italic = True
run.font.color.rgb = GRAY
run.font.name = 'Calibri'
run.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved to: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes ({os.path.getsize(OUTPUT_PATH)/1024:.1f} KB)")
